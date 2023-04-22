#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
-------------------------------------------------------------------------------

  @author: k20045146

   AUTHOR: k20045146

     NAME: analysis_full.py

  PURPOSE: data analysis, Y3 project (FV2) 

QUESTIONS: Do graph theoretic measures depend on overall mean functional connectivity? 
           Which of NR, MR and GSR lead to the best classification? 

-------------------------------------------------------------------------------

    OVERVIEW : 

    ––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––
    - 0 - LIBRARIES 
    ––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––
    - I - DATA INSPECTION 
        - I.1 - Download data 
        - I.2 - Retrieve data 
        - I.3 - time series inspection 
        - I.4 - Check the parcellation process
        - I.5 - Exclude time series corresponding to arbitrary voxels 
        - I.6 - Exclude time series that have values only equal to zero. 
    ––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––
    - II - PARTICIPANT EXCLUSION & ANALYSIS
         - II.1 - Separate the times series of TC and ASD
         - II.2 - Distance matrix for H-O atlas in MNI-152 space
         - II.3 - Before exclusions (+ plots)
         - II.4 - Participant exclusion - unbalanced (+ plots)
         - II.5 - Participant exclusion - balanced FD (+ plots)
         - II.6 - Inducing group differences 
         - II.7 - Mann Whitney U tests - plot
         - II.8 - Linear regressions - plot
    ––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––
    - III - CLASSIFICATION    
          - III.1 - Heterogeneity between scanning sites (+ plot)
          - III.2 - Stratified sampling and diagnosis balancing
          - III.3 - Classification using rbf, poly and linear kSVM
    ––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––
    - IV - DOCUMENTATION & REFERENCES
    ––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––
         
    
        
             
-------------------------------------------------------------------------------   
       


"""

# %%             - 0 - LIBRARIES

# ––– Plots –––
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import StandardScaler
from scipy.spatial import distance
import nibabel as nib
from sklearn.metrics import roc_curve, roc_auc_score, confusion_matrix
import _pickle as cPickle
import re
import ast
from decimal import Decimal
from nilearn.datasets import fetch_abide_pcp
from sklearn.metrics import accuracy_score, recall_score, precision_score
from sklearn.svm import SVC
from collections import defaultdict
from sklearn.model_selection import StratifiedKFold, train_test_split
import statsmodels.api as sm
import pandas as pd
from statsmodels.sandbox.stats.multicomp import multipletests
import scipy.stats as stats
import glob
import bct as bct
import numpy as np
from nilearn import image as nimg
from nilearn import input_data
from nilearn import plotting
import matplotlib.pyplot as plt
import seaborn as sns
import pingouin as pg

# ––– Custom functions (analysis) –––
# Set the directory to where the functions are being stored:
import os
os.chdir('/Users/clarariegis/Desktop/Kings/Y3_2022_23/project/code')
from raincloud import raincloud
from stat_analysis import stat_analysis
from graph_measures import graph_measures

# ––– General –––

# ––– Classification –––


os.chdir('/Users/clarariegis/Downloads')


# %%             - I - PRE-ANALYSIS                        –––––––
# %%%                - I.1 - Dowload data                  |  ✓  |

# Using the 'fetch_abide_pcp()" function to download the ROIs time-series
# into a specified directory. Here, the ROIs parcellated according to the
# Harvard-Oxford atlad that passed the quality control.
# I ran the function twice: once to download with GSR, and once withOUT.
abide = fetch_abide_pcp(data_dir="participants", derivatives=['rois_ho'],
                        global_signal_regression=False, pipeline='ccs',
                        band_pass_filtering=True, quality_checked=True)


# %%%                - I.2 - Retrieve data                 |  ✓  |

# Setting a path to access the GSR data:
path_gsr = 'participants/ABIDE_pcp/ccs/filt_global/*_*_rois_ho.1D'

# ____________ INITIALIZE VARIABLES ____________


all_loaded_id = []   # Store the IDs for all participants.
all_loaded_gsr = []  # Store the GSR time series for all participants.
all_loaded_nr = []   # Store the NR time series for all participants.
# loaded_id_noqc = []  # Quality control False


#     # No quality control
# os.chdir('/Users/clarariegis/Desktop/Kings/Y3_2022_23/project')
# path = 'participants/ABIDE_pcp/ccs/filt_noglobal/*_*_rois_ho.1D'
# for filename in glob.glob(path):

#     # ____________ PCP IDs ____________

#     # Get the ID of the participant whose file is being loaded:
#     loaded_id_noqc += [int(filename.split('_')[-3][2:])]
#     # [-3] -> "-" is to start from the end of the split string. This is because
#     #         the first part of the file names where not always the same.
#     # [2:] -> Not take into account the first two numbers (00).


os.chdir('/Users/clarariegis/Downloads')
for filename in glob.glob(path_gsr):

    # ____________ PCP IDs ____________

    # Get the ID of the participant whose file is being loaded:
    loaded_id = int(filename.split('_')[-3][2:])
    # [-3] -> "-" is to start from the end of the split string. This is because
    #         the first part of the file names where not always the same.
    # [2:] -> Not take into account the first two numbers (00).

    all_loaded_id.append(loaded_id)

    # ____________ STORE GSR & NR time series ____________

    # GSR times series:
    all_loaded_gsr.append(np.loadtxt(filename))

    # NR path - based on the ID of the participant that was loaded
    # from the GSR data:
    path_nr = (filename.split('filt_global/')[0] +
               'filt_noglobal/' +
               filename.split('/')[4])
    # NR times series:
    all_loaded_nr.append(np.loadtxt(path_nr))


# GSR - storing the retrieved IDs and time series in one data frame:
all_pcp_gsr = pd.DataFrame({"ID": all_loaded_id,  # Participants IDs.
                           "ts": all_loaded_gsr})  # Times series corresponding
# to each ID.

# NR - storing the retrieved IDs and time series in one data frame:
all_pcp_nr = pd.DataFrame({"ID": all_loaded_id,   # Participants IDs.
                           "ts": all_loaded_nr})  # Times series corresponding
# to each ID.

# clean variable explorer:
del (all_loaded_id, all_loaded_gsr, all_loaded_nr, filename, loaded_id,
     path_gsr, path_nr)


# %%%                - I.3 - time series inspection        |  ✓  |

# ISSUE: The number of time series (111) in each participants' file (parcellated
# using the Harvard-Oxford atlas) did not match the number of region labels (110).

# Loading the file with the ROIs labels obtained from the ABIDE dataset - it is
# the same on Cyberduck and on the ABIDE webpage.
ho_labels = pd.read_csv("participants/atlas/ho_labels.csv")  # -> 110 labels.

# Checking how many files has 111 time series <=> 111 ROIs.
length = []
for i in range(len(all_pcp_nr["ts"])):
    if all_pcp_gsr["ts"][0].shape[1] == 111:
        length.append(i)

if len(length) == len(all_pcp_gsr["ts"]):
    print("All the participants have 111 ROIs")
# -> it was  true :(


# It also seems that some of the time series only have "0" as values:
zeros_ts = []
for i in range(len(all_pcp_nr["ts"])):

    for j in range(all_pcp_nr["ts"][i].shape[1]):

        if np.all(all_pcp_nr["ts"][i][:, j] == 0):
            # print(i,j)
            zeros_ts.append([i, j])


# %%%                - I.4 - check parcellation process    |  ✓  |


# - Conducting the parcellation on 1 participants to see if the result is the same.

# Load a nifti file for a random participant:
nii_nr_1pcp = nib.load(
    "participants/func_preproc_noglobal/Yale_0050607_func_preproc.nii")


# Load the Harvard-Oxford Atlas:
ho_atlas = nib.load("participants/atlas/ho_roi_atlas.nii")

# Get the image labels. When return_label_names is True, the function returns
# the coordinates + the labels. Here I only stored the labels, hence the "[1]".
ho_img_labels = np.array(plotting.find_parcellation_cut_coords(
    ho_atlas, return_label_names=True)[1])

# Some online documentation pointed out that the "3455" label does not correspond
# to any brain region in the csv label file. Here we check where it is located:
# Index 82 -> consistent with online forum posts.
np.where(ho_img_labels == 3455)
# See Reference section at the end of the code to see the relevant documentation.

# Resampling the atlas to match the dimensions of the nifti file.
ho_resampled = nimg.resample_to_img(
    ho_atlas, nii_nr_1pcp, interpolation='nearest')

# Creating a mask with the H-O atlas.
masker_filt = input_data.NiftiLabelsMasker(labels_img=ho_resampled, detrend=False,
                                           low_pass=0.1, high_pass=0.01, t_r=2)


# Using the mask to average the ts for each ROI
ts_nr_1pcp = masker_filt.fit_transform(nii_nr_1pcp)
# Unsurprisingly the number of ts is the same (111).
print(ts_nr_1pcp.shape[1])

# Storing the time series of the same particpant but pre-parcellated (from ABIDE).
ts_abide_nr_1pcp = all_pcp_nr.loc[all_pcp_nr["ID"] == 50607, "ts"].tolist()[0]

fig, axs = plt.subplots((2))
fig.tight_layout(pad=4)
axs[0].plot(ts_nr_1pcp[:100, 82])
axs[0].set_title("ts 82 (my parcelation)")
axs[1].plot(ts_abide_nr_1pcp[:100, 82])
axs[1].set_title("ts 82 (abide parcelation)")


# %%%                - I.5 - exclude ts causing issue      |  ✓  |


# - Exlude time series 82.

for i in range(len(all_pcp_nr["ts"])):

    all_pcp_nr["ts"][i] = np.delete(all_pcp_nr["ts"][i], 82, axis=1)

    all_pcp_gsr["ts"][i] = np.delete(all_pcp_gsr["ts"][i], 82, axis=1)


# zeros_ts was calculated in section 1.2.b - Time series inspection.
# Let's check again how many ts with only zeros there are after removing the
# ts 82 (which seemed to not match any ROI label).
print(f" There were {len(zeros_ts)} with only 0s.")
zeros_ts = []
for i in range(len(all_pcp_nr["ts"])):

    for j in range(all_pcp_nr["ts"][i].shape[1]):

        if np.all(all_pcp_nr["ts"][i][:, j] == 0):
            # print(i,j)
            zeros_ts.append([i, j])

print(
    f"After removing the ts 82, there are only {len(zeros_ts)} with only zeros left")

# %%%                - I.6 - exclude remaining pcps = 0    |  ✓  |

# - Excluded the participants that have time series with
# only zeros despite having excluded the ts 82 that seemed to not correspond to
# any ROI.

# REMINDER: • so far all the participants are stores in two data frames
# containing there ID and the ts ("all_pcp_nr" = No Regression; "all_pcp_gsr" =
# Global Signal Regression).
# • "zeros_ts" contains the pcp index and the ts index of the ts with only 0s. .
zeros_ts2 = []
for i in range(len(zeros_ts)):
    zeros_ts2 += [zeros_ts[i][0]]

# Just get the single values. This is necessary because some pcps have more than
# one ts causing issue, but we can only excluded each participant once.
zeros_ts2 = list(set(zeros_ts2))

# Excluding the participants:
for i in range(len(zeros_ts2)):
    all_pcp_nr.drop(labels=zeros_ts2[i], axis=0, inplace=True)
    all_pcp_gsr.drop(labels=zeros_ts2[i], axis=0, inplace=True)


# Put the index back to being consecutive numbers - which is not the case since
# some participants were just excluded:
all_pcp_nr.index = np.arange(0, len(all_pcp_nr))
all_pcp_gsr.index = np.arange(0, len(all_pcp_gsr))


# Checking again how many participants have some time series equal to 0s - just
# to make sure none is remaining.
zeros_ts = []
for i in range(len(all_pcp_nr["ts"])):

    for j in range(all_pcp_nr["ts"][i].shape[1]):

        if np.all(all_pcp_nr["ts"][i][:, j] == 0):
            # print(i,j)
            zeros_ts.append([i, j])

print(f"Now {len(zeros_ts)} participant have time series with only zeros.")


# %%            - II - STATISTICAL ANALYSIS                –––––––
# %%%                - II.1 - data preparation             |  ✓  |


# Loading the phenotypic file to know which ID are in the ASD or TC groups.
phenotypic = pd.read_csv(
    'participants/ABIDE_pcp/Phenotypic_V1_0b_preprocessed1.csv')


# Male / females without quality check 
available = phenotypic[phenotypic["subject"].isin(loaded_id_noqc)]
sum(available["SEX"] == 1)
sum(available["SEX"] == 2)
np.min(available["AGE_AT_SCAN"])

# ASD / TC without quality check 
sum(available["DX_GROUP"] == 1)
sum(available["DX_GROUP"] == 2)
np.min(available["AGE_AT_SCAN"])

# ASD / TC with quality check 
abide_qa = phenotypic[phenotypic["subject"].isin(all_pcp_gsr["ID"])]
sum(abide_qa["DX_GROUP"] == 1)
sum(abide_qa["DX_GROUP"] == 2)


# - Separate ASD and TC pcps

# The TC group is categorised as "2", which I switched to "0".
phenotypic['DX_GROUP'][phenotypic['DX_GROUP'] == 2] = 0

# Separate/store the IDs of HC and ASD participants:
id_tc2 = phenotypic.loc[phenotypic['DX_GROUP'] == 0, 'SUB_ID'].tolist()
id_asd2 = phenotypic.loc[phenotypic['DX_GROUP'] == 1, 'SUB_ID'].tolist()


# REMINDER: data for all participants in the NR group was previoulsy stored
# in: "all_pcp_nr" while the GSR group was stored in: "all_pcp_gsr" (see I.2.a).


# Dictionary with NR and GSR time series for TC.
all_pcp_tc = {"nr": all_pcp_nr[all_pcp_nr["ID"].isin(id_tc2)],
              "gsr": all_pcp_gsr[all_pcp_gsr["ID"].isin(id_tc2)]}

# Making sure the index is a list of consecutive numbers.
all_pcp_tc["nr"].index = np.arange(0, len(all_pcp_tc["nr"]))
all_pcp_tc["gsr"].index = np.arange(0, len(all_pcp_tc["gsr"]))


# Dictionary with NR and GSR time series for ASD.
all_pcp_asd = {"nr": all_pcp_nr[all_pcp_nr["ID"].isin(id_asd2)],
               "gsr": all_pcp_gsr[all_pcp_gsr["ID"].isin(id_asd2)]}

# Making sure the index is a list of consecutive numbers.
all_pcp_asd["nr"].index = np.arange(0, len(all_pcp_asd["nr"]))
all_pcp_asd["gsr"].index = np.arange(0, len(all_pcp_asd["gsr"]))


    ## Exclusion
graph_measures(all_pcp_tc["nr"], all_pcp_tc["gsr"], prop_threshold = 0.15, phenotypic=phenotypic, exclusion=True)
graph_measures(all_pcp_asd["nr"], all_pcp_asd["gsr"], prop_threshold = 0.15, phenotypic=phenotypic, exclusion=True)


# %%%                - II.2 - distance matrix              |  ✓  |


# DISTANCE MATRIX
# It will be used to plot the correlation between edges and mFD as a function
# of distance.
# Getting the coordinated of the desikan ROIs in MNI-152 space (excluding the
# background image):
ho_atlas = nib.load("participants/atlas/ho_roi_atlas.nii")
ho_coord = plotting.find_parcellation_cut_coords(ho_atlas, background_label=0)
ho_coord = np.delete(ho_coord, 82, axis=0)  # Delete the "fake" ROI.
# Getting the condensed distance matrix:
cond_dist_mat = distance.pdist(ho_coord, metric="euclidean")


# %%%                - II.3 - before exclusions            |  ✓  |


# Loading the phenotypic file to know which ID are in the ASD or TC groups.
phenotypic = pd.read_csv(
    'participants/ABIDE_pcp/Phenotypic_V1_0b_preprocessed1.csv')
phenotypic = phenotypic.set_index("SUB_ID")



# tc_meas_noexclusion = graph_measures(
#     all_pcp_tc["nr"], all_pcp_tc["gsr"], prop_threshold=0.05, phenotypic=phenotypic, exclusion=False)
# asd_meas_noexclusion = graph_measures(
#     all_pcp_asd["nr"], all_pcp_asd["gsr"], prop_threshold=0.05, phenotypic=phenotypic, exclusion=False)

# mfc_tc = tc_meas_noexclusion["group_nr"]["mean FC"]
# mfc_asd = asd_meas_noexclusion["group_nr"]["mean FC"]


size_plt = 23
size_pcp = 18
col_all = "darkslateblue"
col_asd = "lightcoral"
col_tc = "cornflowerblue"
col_fem = "darkgoldenrod"
col_mal = "wheat"

fig, ax = plt.subplot_mosaic("""
                              CC
                              AB
                              DD
                              EE
                              """, figsize=(13, 27))

plt.subplots_adjust(top=0.7, bottom=0.01, hspace=0.7, wspace=0.3)


# –––––––––  MEAN FC DISTRBUTION –––––––––

# Mean FC distribution for TC group - before exclusion :
ax["A"].hist(mfc_tc, bins=80, color=col_tc)
ax["A"].set_xlabel("mean FC", size=size_plt)          # x label.
ax["A"].set_ylabel("Frequency",  size=size_plt)       # y label.
# set the size of the ticks.
ax["A"].tick_params(axis="both", labelsize=size_plt)
# participants number.
# ax["A"].text(0.62, 16, f"n = {len(mfc_tc)}", size=size_pcp)
# Add median and exclusion boundaries:
ax["A"].axvline(np.median(mfc_tc), color='black', linestyle='-', linewidth=1.4)
upper_boudary = np.median(mfc_tc) + 3 * stats.median_abs_deviation(mfc_tc)
ax["A"].axvline(upper_boudary, color='grey', linestyle='dashed', linewidth=1)
lower_boudary = np.median(mfc_tc) - 3 * stats.median_abs_deviation(mfc_tc)
ax["A"].axvline(lower_boudary, color='grey', linestyle='dashed', linewidth=1)


# Mean FC distribution for ASD patients - before exclusion :
ax["B"].hist(mfc_asd, bins=80, color=col_asd)
ax["B"].set_xlabel("mean FC", size=size_plt)
ax["B"].set_ylabel("Frequency",  size=size_plt)
# set the size of the ticks.
ax["B"].tick_params(axis="both", labelsize=size_plt)
ax["B"].set_ylim(ax["A"].get_ylim())
# ax["B"].text(0.62, 16, f"n = {len(mfc_asd)}", size=size_pcp)
ax["B"].axvline(np.median(mfc_asd), color='black', linestyle='-', linewidth=1.2)
upper_boudary = np.median(mfc_asd) + 3 * stats.median_abs_deviation(mfc_asd)
ax["B"].axvline(upper_boudary, color='black', linestyle='dashed', linewidth=1)
lower_boudary = np.median(mfc_asd) - 3 * stats.median_abs_deviation(mfc_asd)
ax["B"].axvline(lower_boudary, color='black', linestyle='dashed', linewidth=1)


# –––––––––  MEAN FD DISTRBUTION –––––––––

# At this stage only pcp with 0s at ts were exluded (but NONE based on mFD or mFC).
all_pcp_fd = phenotypic.loc[all_pcp_nr["ID"], "func_mean_fd"]
ax["C"].hist(all_pcp_fd, bins=90, color=col_all)
ax["C"].set_xlabel("mean FD",  size=size_plt)
ax["C"].set_ylabel("Frequency",  size=size_plt)
# set the size of the ticks.
ax["C"].tick_params(axis="both", labelsize=size_plt)
# ax["C"].set_title("Distribution of mean FD before exclusion", size = 25, fontweight = "bold")
# ax["C"].text(1.2, 130, f"n = {len(all_pcp_nr)}", size=size_pcp)
ax["C"].axvline(0.4, color='black', linestyle='dashed', linewidth=1)


barWidth = 0.30


all_meas_noexclusion = pd.concat(
    [tc_meas_noexclusion["group_nr"], asd_meas_noexclusion["group_nr"]], axis=0)
all_meas_noexclusion = all_meas_noexclusion.set_index('ID')
rem_pheno = phenotypic.loc[all_meas_noexclusion.index]
rem_ev_noexcl = pd.concat([all_meas_noexclusion, rem_pheno], axis=1)
rem_ev_noexcl["SUB_ID"] = rem_ev_noexcl.index
rem_ev_noexcl.index = np.arange(0, len(rem_ev_noexcl))

TC = rem_ev_noexcl.loc[rem_ev_noexcl["DX_GROUP"] == 2]
ASD = rem_ev_noexcl.loc[rem_ev_noexcl["DX_GROUP"] == 1]

TC_plt = []
ASD_plt = []

for this_site in rem_ev_noexcl["SITE_ID"].unique():

    TC_plt.append(len(TC[TC["SITE_ID"] == this_site]))
    ASD_plt.append(len(ASD[ASD["SITE_ID"] == this_site]))


# Set position of bar on X axis
br0 = np.arange(len(rem_ev_noexcl["SITE_ID"].unique()))
br1 = [x + 0.25 for x in br0]
br2 = [x + barWidth for x in br1]

# Make the plot

ax["D"].bar(br1, TC_plt, color=col_tc, width=barWidth,
            edgecolor='grey', label='TC')
ax["D"].bar(br2, ASD_plt, color=col_asd, width=barWidth,
            edgecolor='grey', label='ASD')


# Adding Xticks
#ax["D"].set_xlabel('Scanning sites', fontweight ='bold', fontsize = 15)
#plt.ylabel('Students passed', fontweight ='bold', fontsize = 15)
ax["D"].set_xticks([r + barWidth for r in range(len(rem_ev_noexcl["SITE_ID"].unique()))],
                   rem_ev_noexcl["SITE_ID"].unique().tolist(), rotation=80)
# set the size of the ticks.
ax["D"].tick_params(axis="both", labelsize=size_plt - 4)
# ax["D"].text(15, 120, f"n = {len(rem_ev_noexcl)}", size=size_pcp)
# ax["D"].text(15, 105, f"n TC = {len(TC)}", size=size_pcp)
# ax["D"].text(15, 90, f"n ASD = {len(ASD)}", size=size_pcp)

# ax["D"].legend()


FEM = rem_ev_noexcl[rem_ev_noexcl["SEX"] == 2]
MAL = rem_ev_noexcl[rem_ev_noexcl["SEX"] == 1]

FEM_plt = []
MAL_plt = []

i = 0
for this_site in rem_ev_noexcl["SITE_ID"].unique():

    FEM_plt.append(len(FEM[FEM["SITE_ID"] == this_site]))
    MAL_plt.append(len(MAL[MAL["SITE_ID"] == this_site]))


# Set position of bar on X axis
br0 = np.arange(len(rem_ev_noexcl["SITE_ID"].unique()))
br1 = [x + 0.25 for x in br0]
br2 = [x + barWidth for x in br1]

# Make the plot

ax["E"].bar(br1, FEM_plt, color=col_fem, width=barWidth,
            edgecolor='grey', label='FEM')
ax["E"].bar(br2, MAL_plt, color=col_mal, width=barWidth,
            edgecolor='grey', label='MAL')


# Adding Xticks
#ax["E"].set_xlabel('Scanning sites', fontweight ='bold', fontsize = 15)
#plt.ylabel('Students passed', fontweight ='bold', fontsize = 15)
ax["E"].set_xticks([r + barWidth for r in range(len(rem_ev_noexcl["SITE_ID"].unique()))],
                   rem_ev_noexcl["SITE_ID"].unique().tolist(), rotation=80)
# set the size of the ticks.
ax["E"].tick_params(axis="both", labelsize=size_plt - 4)
# ax["E"].text(15, 120, f"n = {len(rem_ev_noexcl)}", size=size_pcp)
# ax["E"].text(15, 105, f"n fem = {len(FEM)}", size=size_pcp)
# ax["E"].text(15, 90, f"n mal = {len(MAL)}", size=size_pcp)
# ax["E"].legend()
ax["D"].set_ylim(ax["E"].get_ylim())


import matplotlib.patches as mpatches
all_patch = mpatches.Patch(color = col_all, label= f'All participants (n = {len(all_pcp_nr)})')
asd_patch = mpatches.Patch(color = col_asd, label=f'ASD (n = {len(mfc_asd)})')
tc_patch = mpatches.Patch(color = col_tc, label= f'TC (n = {len(mfc_tc)})')
fem_patch = mpatches.Patch(color = col_fem, label= f'Females (n = {len(FEM)})')
mal_patch = mpatches.Patch(color = col_mal, label= f'Males (n = {len(MAL)})')


fig.legend(handles=[all_patch, asd_patch, tc_patch, fem_patch, mal_patch], loc = (0.63,0.813), fontsize = size_plt)

ax["C"].set_title("Mean FD distribution", size = size_plt)
ax["A"].text(x = 0.5, y = 25, s = "Mean FC distribution in TC and ASD", size = size_plt)
ax["D"].set_title("ASD and TC distribution across scanning sites", size = size_plt)
ax["E"].set_title("Female and male distribution across scanning sites", size = size_plt)


plt.show()


# %%%                - II.4 - FD & FD 0.4 - 0.1            |  ✓  |


size_plt = 20
plt_size = 20
size_pcp = 13
col_all = "darkslateblue"
col_asd = "lightcoral"
col_tc = "cornflowerblue"
n_round = 3

plt.rcParams.update({"axes.grid" : False})

list_fd_max = [0.4]  # , 0.3, 0.2, 0.1
# , "maximum mean FD : 0.3", "maximum mean FD : 0.2", "maximum mean FD : 0.1"
group_labels = ["max mean FD: 0.1"] # ["maximum mean FD : 0.4"] #, "balanced mean FD"]

with open(r"clean_unbalanced_dat.pickle", "rb") as input_file:
    rem_everything_un2 = cPickle.load(input_file)

with open(r"clean_balanced_dat.pickle", "rb") as input_file:
    rem_everything_bal2 = cPickle.load(input_file)

rem_everything_fd01 = rem_everything_un2[rem_everything_un2["func_mean_fd"] <= 0.1]

data_2 = [rem_everything_fd01] # [rem_everything_un2] #, rem_everything_bal2]

for (this_label, this_dat) in zip(group_labels, data_2):
    col_all = "darkslateblue"
    phenotypic = pd.read_csv('participants/ABIDE_pcp/Phenotypic_V1_0b_preprocessed1.csv')

    #     # TC outliers:
    # funct_measures_tc = graph_measures(all_pcp_tc["nr"], all_pcp_tc["gsr"], prop_threshold = 0.15, phenotypic = phenotypic, fd_lim = this_fd_max, exclusion = True)

    # # Number of participants after non finite exclusion: 454
    # # Number of mFC outliers excluded: 22
    # # Number of mFD outliers excluded: 6

    #     # ASD outliers:
    # funct_measures_asd = graph_measures(all_pcp_asd["nr"], all_pcp_asd["gsr"], prop_threshold = 0.15, phenotypic = phenotypic, fd_lim = this_fd_max, exclusion = True)

    # # Number of participants after non finite exclusion: 390
    # # Number of mFC outliers excluded: 23
    # # Number of mFD outliers excluded: 10

    rem_pheno_tc = this_dat.loc[this_dat["DX_GROUP"] == 0]
    rem_pheno_asd = this_dat.loc[this_dat["DX_GROUP"] == 1]

    # –––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––
    # –––––––––––––––––– FIGURE 1 : –––––––––––––––––––––––––––––––––––––––––––
    # –––––––––––––––––– FRAMEWISE DISPLACEMENT AND FUNCTIONAL CONNECTIVITY –––
    # –––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––

    # ________ Framewise Displacement Raincloud ________

    # Compare the framewise displacement of both groups
    # just putting the FD in a variable to make it more handy in the mann whitney U
    rem_fd_tc = rem_pheno_tc["func_mean_fd"]
    rem_fd_asd = rem_pheno_asd["func_mean_fd"]
    # len(rem_fd_tc) # checking the the number of participants.
    # len(rem_fd_asd)
    
    # assessing whether the difference is significant.
    mann_fd = stats.mannwhitneyu(rem_fd_tc, rem_fd_asd)
    rbc = float(pg.mwu(rem_fd_tc, rem_fd_asd)["RBC"])
    # Reporting a few descriptive stats and the p-value.
    # print(f"""mean FD TC: {np.median(rem_fd_tc)} and ASD: {np.median(rem_fd_asd)}; 
    # median absolute deviation TC: {stats.median_abs_deviation(rem_fd_tc)} and ASD : {stats.median_abs_deviation(rem_fd_asd)} 
    # p-value: {mann_fd[1]} """)

    rem_fd = (rem_fd_tc, rem_fd_asd)

    fig, ax = plt.subplots(4, 3, figsize=(15, 14))
    plt.subplots_adjust(top=0.8, bottom=0.01, hspace=0.5, wspace=0.4)
    plt.grid(False)
    #plt.tick_params(axis = "both" , labelsize=size_plt)

    # Add a general title to the figure (specifying which mean FD cutoff is used).
    fig.suptitle(this_label,  fontsize=20, fontweight="bold", y=0.86)

    axx = ax[0, 0]

    p_value = mann_fd[1]
    if p_value < 0.01:
        p_value = f"= {Decimal(p_value):.2e}"
    else:
        p_value = f" = {round(p_value, 3)}"
    
    title = r"r$_{rb}$ = " + f"{round(rbc, n_round)}, p {p_value}"
    raincloud(rem_fd, feature1="TC", feature2="ASD",  x_label="Mean FD",
              title= title, colors="on", ax=axx, size=size_plt)

    # ________ Famewise Displacement Histogram ________

    # FD distribution after balancing : ASD
    ax[1, 0].hist(rem_pheno_asd["func_mean_fd"], bins=80, color=col_asd)
    ax[1, 0].set_xlabel("Mean FD", fontsize=size_plt)
    ax[1, 0].set_ylabel("Frequency", fontsize=size_plt)
    # ax[1, 0].text(0.8, 0.9, f'n = {len(rem_pheno_asd["func_mean_fd"])}', horizontalalignment='center',
    #               verticalalignment='center', fontsize=size_pcp, transform=ax[1, 0].transAxes)

    # FD distribution after balancing : TC
    ax[2, 0].hist(rem_pheno_tc["func_mean_fd"], bins=80, color=col_tc)
    ax[2, 0].set_xlabel("Mean FD", fontsize=size_plt)
    ax[2, 0].set_ylabel("Frequency", fontsize=size_plt)
    # ax[2, 0].text(0.8, 0.9, f'n = {len(rem_pheno_tc["func_mean_fd"])}', horizontalalignment='center',
    #               verticalalignment='center', fontsize=size_pcp, transform=ax[2, 0].transAxes)
    ax[1, 0].set_ylim(ax[2, 0].get_ylim())

    # FD distribution after balancing : ALL
    all_mFD = rem_pheno_tc["func_mean_fd"].tolist(
    ) + rem_pheno_asd["func_mean_fd"].tolist()
    ax[3, 0].hist(all_mFD, bins=80, color=col_all)
    ax[3, 0].set_xlabel("Mean FD", fontsize=size_plt)
    ax[3, 0].set_ylabel("Frequency", fontsize=size_plt)
    # ax[3, 0].text(0.8, 0.9, f'n = {len(all_mFD)}', horizontalalignment='center',
    #               verticalalignment='center', fontsize=size_pcp, transform=ax[3, 0].transAxes)

    # ________ Functional Connectivity Raincloud ________
    axx = ax[0, 1]

    mann_whi = stats.mannwhitneyu(rem_pheno_tc["mean FC"], rem_pheno_asd["mean FC"])
    u, p = mann_whi
    # makes the titles cleaner on the plot (removes the "Name" and "dtype")
    p_value = float(p)
    
    # Rank biserial correlation:
    rbc = float(pg.mwu(rem_pheno_tc["mean FC"], rem_pheno_asd["mean FC"])["RBC"])
    title = r"r$_{rb}$ = " + f"{round(rbc, n_round)}, p = {round(p_value, n_round)}"
    
    mFC = [rem_pheno_tc["mean FC"], rem_pheno_asd["mean FC"]]
    raincloud(mFC, feature1="TC", feature2="ASD", x_label="Mean FC",
              title= title, colors="on", ax=axx, size=size_plt)

    # ________ Functional Connectivity Histogram ________

    #ax.get_shared_y_axes().join(ax[1,1], ax[2,1], ax[3,1])

    # FC distribution : ASD
    ax[1, 1].hist(rem_pheno_asd["mean FC"], bins=80, color=col_asd)
    ax[1, 1].set_xlabel("Mean FC", fontsize=size_plt)
    ax[1, 1].set_ylabel("Frequency", fontsize=size_plt)
    # ax[1, 1].text(0.8, 0.9, f'n = {len(rem_pheno_asd["mean FC"])}', horizontalalignment='center',
    #               verticalalignment='center', fontsize=size_pcp, transform=ax[1, 1].transAxes)

    # FC distribution : TC
    ax[2, 1].hist(rem_pheno_tc["mean FC"], bins=80, color=col_tc)
    ax[2, 1].set_xlabel("Mean FC", fontsize=size_plt)
    ax[2, 1].set_ylabel("Frequency", fontsize=size_plt)
    # ax[2, 1].text(0.8, 0.9, f'n = {len(rem_pheno_tc["mean FC"])}', horizontalalignment='center',
    #               verticalalignment='center', fontsize=size_pcp, transform=ax[2, 1].transAxes)
    ax[1, 1].set_ylim(ax[2, 1].get_ylim())

    # FC distribution : ALL
    # convert into list because it was in pandas core series so it wasn't working.
    all_mFC = rem_pheno_tc["mean FC"].tolist(
    ) + rem_pheno_asd["mean FC"].tolist()
    ax[3, 1].hist(all_mFC, bins=80, color=col_all)
    ax[3, 1].set_xlabel("Mean FC", fontsize=size_plt)
    ax[3, 1].set_ylabel("Frequency", fontsize=size_plt)
    # ax[3, 1].text(0.8, 0.9, f'n = {len(all_mFC)}', horizontalalignment='center',
    #               verticalalignment='center', fontsize=size_pcp, transform=ax[3, 1].transAxes)

    # ________ Correlation FD and FC ________
    # ASD
    sns.regplot(x=rem_pheno_asd["func_mean_fd"], y=mFC[1], scatter_kws={
                's': 8}, ax=ax[1, 2], color=col_asd)
    ax[1, 2].set_xlabel("Mean FD", fontsize=size_plt)
    ax[1, 2].set_ylabel("Mean FC", fontsize=size_plt)
    # ax[1, 2].text(0.8, 0.9, f'n = {len(mFC[1])}', horizontalalignment='center',
    #               fontsize=size_pcp, verticalalignment='center', transform=ax[1, 2].transAxes)
    r, p = stats.pearsonr(x=rem_pheno_asd["func_mean_fd"], y=mFC[1])
    # round them cause we don't need a tone of decimals
    ax[1, 2].set_title(
        f"r = {round(r, 3)}, p = {round(p, 3)}", fontsize=size_plt)

    if p < 0.01:
        p = f"= {Decimal(p):.2e}"
    else:
        p = f" = {round(p, 3)}"

    ax[1, 2].set_title(f"r = {round(r, 3)}, p {p}", fontsize=size_plt)

    # TC
    sns.regplot(x=rem_pheno_tc["func_mean_fd"], y=mFC[0], scatter_kws={
                's': 8}, ax=ax[2, 2], color=col_tc)
    ax[2, 2].set_xlabel("Mean FD", fontsize=size_plt)
    ax[2, 2].set_ylabel("Mean FC", fontsize=size_plt)
    # ax[2, 2].text(0.8, 0.9, f'n = {len(mFC[0])}', horizontalalignment='center',
    #               verticalalignment='center', fontsize=size_pcp, transform=ax[2, 2].transAxes)
    r, p = stats.pearsonr(x=rem_pheno_tc["func_mean_fd"], y=mFC[0])

    if p < 0.01:
        p = f"= {Decimal(p):.2e}"
    else:
        p = f" = {round(p, 3)}"

    ax[2, 2].set_title(f"r = {round(r, 3)}, p {p}", fontsize=size_plt)

    # ALL
    sns.regplot(x=all_mFD, y=all_mFC, scatter_kws={
                's': 8}, ax=ax[3, 2], color=col_all)
    ax[3, 2].set_xlabel("Mean FD", fontsize=size_plt)
    ax[3, 2].set_ylabel("Mean FC", fontsize=size_plt)
    # ax[3, 2].text(0.8, 0.9, f'n = {len(all_mFC)}', horizontalalignment='center',
    #               verticalalignment='center', fontsize=size_pcp, transform=ax[3, 2].transAxes)
    r, p = stats.pearsonr(x=all_mFD, y=all_mFC)

    if p < 0.01:
        p = f"= {Decimal(p):.2e}"
    else:
        p = f" = {round(p, 3)}"

    ax[3, 2].set_title(f"r = {round(r, 3)}, p {p}", fontsize=size_plt)

    # ________ Correlation FD and FC as a function of distance ________

    # FD from all remaining participants.
    rem_pcp_fd = this_dat["func_mean_fd"]

    # Vectorised upper triangular of remaining participants:
    rem_vec_all = np.array(this_dat["vectorised matrix"].values.tolist())

    # Calculate the correlation between edges and FD accross participants.
    corr_fd_fc = np.zeros(rem_vec_all.shape[1])
    for i in range(rem_vec_all.shape[1]):
        corr_fd_fc[i] = stats.pearsonr(rem_vec_all[:, i], rem_pcp_fd)[0]

    # Put the correlation coefficients and the distances in one dataframe.
    corr_dist = pd.DataFrame({"corr_fd_fc": corr_fd_fc, "distance": cond_dist_mat})
    corr_dist = corr_dist.dropna()  # Get rid of the rows that have NaN values.
    
    values = np.vstack([corr_dist["corr_fd_fc"], corr_dist["distance"]])
    kernel = stats.gaussian_kde(values)(values)

    sns.scatterplot(corr_dist,
                    x="distance",
                    y="corr_fd_fc",
                    c=kernel,
                    cmap="viridis",
                    ax=ax[0, 2])

    sns.regplot(data=corr_dist, x="distance", y="corr_fd_fc",
                scatter=False, ax=ax[0, 2], color="red")
    
    ax[0, 2].set_xlabel("Distance", fontsize=size_plt)
    ax[0, 2].set_ylabel("Corr mFD and mFC", fontsize=size_plt)
    
    r, p = stats.pearsonr(x=corr_dist["distance"], y=corr_dist["corr_fd_fc"])

    if p < 0.01:
        p = f"= {Decimal(p):.2e}"
    else:
        p = f" = {round(p, 3)}"

    # round them cause we don't need a tone of decimals
    ax[0, 2].set_title(f"r = {round(r, 5)}, p {p}", fontsize=size_plt)
    
    from matplotlib.lines import Line2D
    red_line = Line2D([0], [0], linewidth=3,  label=f'ASD (n = {len(rem_pheno_asd["mean FC"])})', color='lightcoral')
    blue_line = Line2D([0], [0], linewidth=3,  label=f'TC (n = {len(rem_pheno_tc["mean FC"])})', color='cornflowerblue')
    purple_line =  Line2D([0], [0], linewidth=3,  label=f'All participants (n = {len(all_mFC)})', color = col_all)
    
    fig.legend(handles=[red_line, blue_line, purple_line], fontsize = plt_size)

    # –––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––
    # –––––––––––––––––– FIGURE 2 : –––––––––––––––––––––––––––––––––––––––––––
    # –––––––––––––––––– EDGES WEIGHTS FOR NR, MR and GSR –––––––––––––––––––––
    # –––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––

    col_all = "darkslateblue"
    id_asd = rem_pheno_asd["subject"]
    id_tc = rem_pheno_tc["subject"]
    id_all = pd.concat([id_asd , id_tc])
    
    ts_all_gsr = all_pcp_gsr[all_pcp_gsr["ID"].isin(id_all)]
    ts_all_nr = all_pcp_nr[all_pcp_nr["ID"].isin(id_all)]
    ts_all_gsr.index = np.arange(len(ts_all_gsr))
    ts_all_nr.index = np.arange(len(ts_all_nr))

    all_measures = graph_measures(ts_all_nr, ts_all_gsr, prop_threshold = 0.15, phenotypic=phenotypic, exclusion=False)

    # Average the vectorised upper triangulars of the connectivity matrices
    # for all three conditions (NR, MR and GSR).
    vec_nr = all_measures["group_nr"]["vectorised matrix"].tolist()
    vec_gsr = all_measures["group_gsr"]["vectorised matrix"].tolist() 
    vec_mreg = all_measures["group_mreg"]["vectorised matrix"].tolist() 


    mfc_nr = pd.DataFrame(np.mean(vec_nr, axis=0))
    mfc_gsr = pd.DataFrame(np.mean(vec_gsr, axis=0))
    mfc_mreg = pd.DataFrame(np.mean(vec_mreg, axis=0))


    # Create a figure with six subplots.
    fig2, ax2 = plt.subplots(2, 3, figsize=(14, 10))
    fig2.subplots_adjust(top=0.6, bottom=0.01, hspace=0.5, wspace=0.3)
    # Add a general title to the figure.
    fig2.suptitle(this_label,  fontsize=20, fontweight="bold", y=0.73)
    # Number of participants
    fig2.text(x=0.2, y=0.65,  s=f"n = {len(vec_gsr)}", size = size_plt )
    plt.grid(False)
    
    plt_size = 20
    
    # GSR in function of NR:
    ax2[0, 0].scatter(x=mfc_nr, y=mfc_gsr, s=2, c=col_all)
    ax2[0, 0].set_xlabel("NR edge weight", fontsize=plt_size)
    ax2[0, 0].set_ylabel("GSR edge weight", fontsize=plt_size)

    # MR in function of NR:
    ax2[0, 1].scatter(x=mfc_nr, y=mfc_mreg, s=2, c=col_all)
    ax2[0, 1].set_xlabel("NR edge weight", fontsize= plt_size)  # averaged weights
    ax2[0, 1].set_ylabel("MR edge weight ", fontsize= plt_size)  # averaged weights

    # MR in function of GSR:
    ax2[0, 2].scatter(x=mfc_gsr, y=mfc_mreg, s=2, c=col_all)
    ax2[0, 2].set_xlabel("GSR edge weight", fontsize=plt_size)
    ax2[0, 2].set_ylabel("MR edge weight", fontsize=plt_size)

    # ____ edge distribution ____

    # NR 
        # averaged edges
    mfc_nr.plot(kind = "kde", ax = ax2[1, 0], legend=False, color = col_all)
    
        # edges of pcp with max mean FC
    meas = all_measures["group_nr"]
    maxi = pd.DataFrame(meas["vectorised matrix"][meas['mean FC'].idxmax()])
    maxi.plot(kind = "kde", ax = ax2[1, 0], legend=False, color = "red")
    
        # edges of pcp with min mean FC
    mini = pd.DataFrame(meas["vectorised matrix"][meas['mean FC'].idxmin()])
    mini.plot(kind = "kde", ax = ax2[1, 0], legend=False, color = "blue")
     
    
    # MR
    mfc_mreg.plot(kind = "kde", ax = ax2[1, 1], legend=False , color = col_all)

        # edges of pcp with max mean FC
    meas = all_measures["group_mreg"]
    maxi = pd.DataFrame(meas["vectorised matrix"][meas['mean FC'].idxmax()])
    maxi.plot(kind = "kde", ax = ax2[1, 1], legend=False, color = "red")
    
        # edges of pcp with min mean FC
    mini = pd.DataFrame(meas["vectorised matrix"][meas['mean FC'].idxmin()])
    mini.plot(kind = "kde", ax = ax2[1, 1], legend=False, color = "blue")


    # MR
        # averaged edges
    mfc_gsr.plot(kind = "kde", ax = ax2[1, 2], legend=False , color = col_all)

        # edges of pcp with max mean FC
    meas = all_measures["group_gsr"]
    maxi = pd.DataFrame(meas["vectorised matrix"][meas['mean FC'].idxmax()])
    maxi.plot(kind = "kde", ax = ax2[1, 2], legend=False, color = "red")
    
        # edges of pcp with min mean FC
    mini = pd.DataFrame(meas["vectorised matrix"][meas['mean FC'].idxmin()])
    mini.plot(kind = "kde", ax = ax2[1, 2], legend=False, color = "blue")

    ax2[1, 0].set_ylim(ax2[1, 2].get_ylim())
    ax2[1, 1].set_ylim(ax2[1, 2].get_ylim())

    ax2[1, 0].set_xlabel("NR edge weight", fontsize=plt_size)
    ax2[1, 1].set_xlabel("MR edge weight", fontsize=plt_size)
    ax2[1, 2].set_xlabel("GSR edge weight", fontsize=plt_size)
    
    ax2[1, 0].set_ylabel("Density", fontsize=plt_size)
    ax2[1, 1].set_ylabel("Density", fontsize=plt_size)
    ax2[1, 2].set_ylabel("Density", fontsize=plt_size)
    
    red_line = Line2D([0], [0], linewidth=3,  label='Max mean FC', color='red')
    blue_line = Line2D([0], [0], linewidth=3,  label='Min mean FC', color='blue')
    purple_line =  Line2D([0], [0], linewidth=3,  label=f'Average (n = {len(meas)})', color= col_all)
    
    fig2.legend(handles=[red_line, blue_line, purple_line], fontsize = plt_size)

    plt.show()
    
    
    # –––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––
    # –––––––––––––––––– FIGURE 3 & 4 : –––––––––––––––––––––––––––––––––––––––
    # –––––––––––––––––– STATISTICAL ANALYSIS –––––––––––––––––––––––––––––––––
    # –––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––
    
    # Corresponding NR time series.
    ts_asd_nr = all_pcp_nr[all_pcp_nr["ID"].isin(id_asd)]
    # Corresponding GSR ts.
    ts_asd_gsr = all_pcp_gsr[all_pcp_gsr["ID"].isin(id_asd)]
    ts_asd_nr.index = np.arange(len(ts_asd_nr))
    ts_asd_gsr.index = np.arange(len(ts_asd_gsr))

    # Corresponding NR time series.
    ts_tc_nr = all_pcp_nr[all_pcp_nr["ID"].isin(id_tc)]
    # Corresponding GSR ts.
    ts_tc_gsr = all_pcp_gsr[all_pcp_gsr["ID"].isin(id_tc)]
    ts_tc_nr.index = np.arange(len(ts_tc_nr))
    ts_tc_gsr.index = np.arange(len(ts_tc_gsr))

    plt_size = 25
    
    two_fig = ["mann", "linreg"]

    for figure in two_fig:


        ii = np.arange(start=0, stop=3, step=1)  # Range of values of eta.
        jj = np.arange(start=0, stop=9, step=1)  # Range of values of gamma.

        # Repeating the ranges, so that we'll get every possible combination of eta-gamma
        # there are XXXX values of each, so we end up with an array of XXXX values
        iis = np.tile(ii, len(jj))  # ([XXXXX]).
        jjs = np.repeat(jj, len(ii))  # ([XXXXX])].
        axes = list(zip(jjs, iis))

        list_plots = ["cluster", "charpath", "assort"]
        groups = ["group_nr", "group_mreg", "group_gsr"]
        thresholds = [0.05, 0.15, 0.25]

        fig, ax = plt.subplots(9, 3, figsize=(20, 35))
        plt.subplots_adjust(top=0.7, bottom=0.02, hspace=0.55, wspace=0.2)

        y=0.73
        fig.text(s="GC ",  x=0.28,  y=y, fontsize=25)
        fig.text(s="CPL",  x=0.55,  y=y, fontsize=25)
        fig.text(s="A",  x=0.83,  y=y, fontsize=25)

        x = 0.085
        fig.text(s="5% ",  x=x,  y=0.67, rotation=90, fontsize=25)
        fig.text(s="15% ",  x=x,  y=0.59, rotation=90, fontsize=25)
        fig.text(s="25%",  x=x,  y=0.51, rotation=90, fontsize=25)
        fig.text(s="NR",  x=0.02,  y=0.593, rotation=90, fontsize=25)

        fig.text(s="5% ",  x=x,  y=0.43, rotation=90, fontsize=25)
        fig.text(s="15% ",  x=x,  y=0.352, rotation=90, fontsize=25)
        fig.text(s="25% ",  x=x,  y=0.274, rotation=90, fontsize=25)
        fig.text(s="MR",  x=0.02,  y=0.355, rotation=90, fontsize=25)

        fig.text(s="5%",  x=x,  y=0.198, rotation=90, fontsize=25)
        fig.text(s="15%",  x=x,  y=0.12, rotation=90, fontsize=25)
        fig.text(s="25%",  x=x,  y=0.04, rotation=90, fontsize=25)
        fig.text(s="GSR",  x=0.02,  y=0.12, rotation=90, fontsize=25)
        
        # fig.suptitle(this_label,  fontsize=35, fontweight="bold", y=0.78)

    # –––––––––––––––––

        i = 0

        size_plt = 20

        for this_group in groups:
            
                for this_thresh in thresholds:
            
                    if figure == "mann":
        
                        tc_measures = graph_measures(ts_tc_nr, ts_tc_gsr, prop_threshold=this_thresh, phenotypic=phenotypic, exclusion=False)
        
                        asd_measures = graph_measures(ts_asd_nr, ts_asd_gsr, prop_threshold=this_thresh, phenotypic=phenotypic, exclusion=False)
        
                    else: 
                        
                        all_measures = graph_measures(ts_all_nr, ts_all_gsr, prop_threshold = this_thresh, phenotypic=phenotypic, exclusion=False)
        
        
                    for this_plot in list_plots:
    
                        if figure == "mann":
    
                            asd = tc_measures[this_group][this_plot].tolist()
    
                            tc = asd_measures[this_group][this_plot].tolist()
    
                            label = " "
    
                            data = [tc, asd]
    
                            mann_whi = stats.mannwhitneyu(tc, asd)
                            rbc = float(pg.mwu(tc, asd)["RBC"])
                            
                            p_value = mann_whi[1]
    
                            col = "on"
    
                            if p_value > 0.1:
                                col = "off"
    
                            if p_value < 0.01:
                                p_value = f"= {Decimal(p_value):.2e}"  # "< 0.01"
    
                            else:
                                p_value = f" = {round(p_value, 3)}"
    
                            axx = axes[i]
                            
                            title = r"r$_{rb}$ = " + f"{round(rbc, n_round)}, p {p_value}"
                            
                            raincloud(data_x=data, feature1="TC", feature2="ASD",  x_label=label,
                                      title = title, colors=col, ax=ax[axx], size=size_plt)
    
                        else:
                            all_fc = all_measures[this_group]["mean FC"]
                            all_meas = all_measures[this_group][this_plot]
    
                            label = " "  # "Global clustering"
    
                            pearson = stats.pearsonr(all_fc, all_meas)
    
                            r_value = pearson[0]
                            p_value = pearson[1]
    
                            if p_value > 0.05:
                                col_all = "darkgrey"
                                
                            else: col_all = "darkslateblue"
    
                            if p_value < 0.01:
    
                                p_value = f"= {Decimal(p_value):.2e}"  # " < 0.01"
    
                            else:
                                p_value = f" = {round(p_value, 3)}"
    
                            axx = axes[i]
    
                            sns.regplot(x=all_fc, y=all_meas, scatter_kws={'s': 5},
                                        color=col_all, ax=ax[axx]).set_title(label=f"r = {round(r_value, 3)}, p {p_value}", fontsize = size_plt)
    
                            ax[axx].tick_params(axis="both", labelsize=20)
                            ax[axx].set_xlabel(label)
                            ax[axx].set_ylabel(label)
    
                        i += 1
    
        plt.show()


# %%%                - II.5 - exclusion - balanced FD      |  ✓  |


# NR measures were generated priviously without balancing mean FD. It was
# stored under "rem_everything_un", where contains the graph measures and
# phenotypic data after participant exclusion for FD exceeding 0.4-0.1.


# rem_everything_un["0.4"].to_csv("clean_unbalanced_dat.csv", index = True)

# rem_everything_un["0.4"].to_pickle('clean_unbalanced_dat.df')


with open(r"clean_unbalanced_dat.pickle", "wb") as output_file:
    cPickle.dump(rem_everything_un["0.4"], output_file)

with open(r"clean_unbalanced_dat.pickle", "rb") as input_file:
    rem_everything_un2 = cPickle.load(input_file)

# Extracting TC participant:
tc_everyt_fd04 = rem_everything_un2.loc[rem_everything_un2["DX_GROUP"] == 0]
# Put the IDs back into their original column.
tc_everyt_fd04["SUB_ID"] = tc_everyt_fd04.index
# Make the index a list of consecutive numbers.
tc_everyt_fd04.index = np.arange(0, len(tc_everyt_fd04))

# Extracting ASD participant:
asd_everyt_fd04 = rem_everything_un2.loc[rem_everything_un2["DX_GROUP"] == 1]
# Put the IDs back into their original column.
asd_everyt_fd04["SUB_ID"] = asd_everyt_fd04.index
# Make the index a list of consecutive numbers.
asd_everyt_fd04.index = np.arange(0, len(asd_everyt_fd04))


i = 0
p = 0

asd_everyt_fd04_bal = asd_everyt_fd04.copy()

while p < 0.05:  # while there is a significant difference in mFD between ASD and TC...

    # ... exclude ASD participants that have the highest mFD.
    asd_everyt_fd04_bal = asd_everyt_fd04_bal.drop(
        asd_everyt_fd04_bal['func_mean_fd'].idxmax())

    # Run a Mann Whitney U test to see if the difference is still significant.
    mann_whi = stats.mannwhitneyu(asd_everyt_fd04_bal['func_mean_fd'],  # ASD FD
                                  tc_everyt_fd04['func_mean_fd'])  # TC FD

    # Store the p-value to assess whether it is still below the 0.05 threshold.
    p = mann_whi[1]

    i += 1

    print(p)


# Dataframe containing the NR measures and the phenotypic information of the
# participants remaining the balanced FD group.
rem_everyt_bal = pd.concat((tc_everyt_fd04, asd_everyt_fd04_bal))
print(len(rem_everyt_bal))

with open(r"clean_balanced_dat.pickle", "wb") as output_file:
    cPickle.dump(rem_everyt_bal, output_file)

with open(r"clean_balanced_dat.pickle", "rb") as input_file:
    rem_everything_bal2 = cPickle.load(input_file)


fig, ax = plt.subplots(1, 2, figsize=(13, 5))
plt.subplots_adjust(top=0.8, bottom=0.01, hspace=0.5, wspace=0.3)
i = 0
list_plots = ["func_mean_fd", "mean FC"]

for this_plot in list_plots:

    asd = asd_everyt_fd04_bal[this_plot]

    tc = tc_everyt_fd04[this_plot]

    label = "mean FD"

    if this_plot == "vectorised matrix":

        asd = asd_everyt_fd04_bal[this_plot]

        tc = tc_everyt_fd04[this_plot]

        label = "mean FC"

    data = [tc, asd]

    mann_whi = stats.mannwhitneyu(tc, asd)

    p_value = mann_whi[1]

    if p_value < 0.01:
        p_value = "< 0.01"

    else:
        p_value = f" = {round(p_value, 3)}"

    raincloud(data, feature1="TC", feature2="ASD",  x_label=label,
              title=f"p {p_value}", colors="on", ax=ax[i], size=size_plt)

    i += 1


#%%%

all_measures = graph_measures(ts_asd_nr, ts_asd_gsr, prop_threshold = 0.05, phenotypic=phenotypic, exclusion=False)

rem_pheno_tc = rem_everything_un2.loc[rem_everything_un2["DX_GROUP"] == 0]
rem_pheno_asd = rem_everything_un2.loc[rem_everything_un2["DX_GROUP"] == 1]

rem_pheno_tc = rem_everything_bal2.loc[rem_everything_bal2["DX_GROUP"] == 0]
rem_pheno_asd = rem_everything_bal2.loc[rem_everything_bal2["DX_GROUP"] == 1]

np.median(rem_pheno_tc["cluster"])
stats.median_abs_deviation(rem_pheno_tc["cluster"])
stats.mannwhitneyu(x = rem_pheno_tc["cluster"], y = rem_pheno_asd["cluster"])

np.median(rem_pheno_asd["cluster"])
stats.median_abs_deviation(rem_pheno_asd["cluster"])
stats.mannwhitneyu(x = rem_pheno_tc["cluster"], y = rem_pheno_asd["cluster"])


# %%%                - II.6 - inducing group differences   |  ✓  | to comment



# # Extracting TC participant:
# tc_everyt_fd04 = rem_everything_un2.loc[rem_everything_un2["DX_GROUP"] == 0]
# # Put the IDs back into their original column.
# tc_everyt_fd04["SUB_ID"] = tc_everyt_fd04.index
# # Make the index a list of consecutive numbers.
# tc_everyt_fd04.index = np.arange(0, len(tc_everyt_fd04))

# # Extracting ASD participant:
# asd_everyt_fd04 = rem_everything_un2.loc[rem_everything_un2["DX_GROUP"] == 1]
# # Put the IDs back into their original column.
# asd_everyt_fd04["SUB_ID"] = asd_everyt_fd04.index
# # Make the index a list of consecutive numbers.
# asd_everyt_fd04.index = np.arange(0, len(asd_everyt_fd04))


# Extracting TC participant:
tc_everyt_fd04 = rem_everything_bal2.loc[rem_everything_bal2["DX_GROUP"] == 0]
# Put the IDs back into their original column.
tc_everyt_fd04["SUB_ID"] = tc_everyt_fd04.index
# Make the index a list of consecutive numbers.
tc_everyt_fd04.index = np.arange(0, len(tc_everyt_fd04))

# Extracting ASD participant:
asd_everyt_fd04 = rem_everything_bal2.loc[rem_everything_bal2["DX_GROUP"] == 1]
# Put the IDs back into their original column.
asd_everyt_fd04["SUB_ID"] = asd_everyt_fd04.index
# Make the index a list of consecutive numbers.
asd_everyt_fd04.index = np.arange(0, len(asd_everyt_fd04))



size_plt = 35

fig, ax = plt.subplots(1, 4, figsize=(40, 5))
plt.subplots_adjust(top=0.8, bottom=0.01, hspace=0.5, wspace=0.3)
i = 0
list_plots = ["mean FC", "cluster", "charpath", "assort"]
labels = ["mean FC", "GC", "CPL", "A"]

for this_plot, label in zip(list_plots, labels):

    asd = asd_everyt_fd04[this_plot]

    tc = tc_everyt_fd04[this_plot]

    data = [tc, asd]

    mann_whi = stats.mannwhitneyu(tc, asd)

    u, p = mann_whi
    
    rbc = float(pg.mwu(tc, asd)["RBC"])

    if p < 0.01:
        p = f"= {Decimal(p):.2e}"

    else:
        p = f" = {round(p, 3)}"

        # U = {round(u, 3)},
        
    title = r"r$_{rb}$ = " + f"{round(rbc, n_round)}, p {p}"
    
    raincloud(data, feature1="TC", feature2="ASD",  x_label=label,
              title = title, colors="on", ax=ax[i], size=size_plt)

    i += 1


asd_everyt_high = asd_everyt_fd04.nlargest(
    int(len(asd_everyt_fd04) * 0.75), columns="mean FC")
tc_everyt_low = tc_everyt_fd04.nsmallest(
    int(len(tc_everyt_fd04) * 0.75), columns="mean FC")


fig, ax = plt.subplots(1, 4, figsize=(40, 5))
plt.subplots_adjust(top=0.8, bottom=0.01, hspace=0.5, wspace=0.3)
i = 0
list_plots = ["mean FC", "cluster", "charpath", "assort"]
labels = ["mean FC", "GC", "CPL", "A"]


for this_plot, label in zip(list_plots, labels):

    asd = asd_everyt_high[this_plot]

    tc = tc_everyt_low[this_plot]

    data = [tc, asd]

    mann_whi = stats.mannwhitneyu(tc, asd)

    u, p = mann_whi
    rbc = float(pg.mwu(tc, asd)["RBC"])

    if p < 0.01:
        p = f"= {Decimal(p):.2e}"

    else:
        p = f" = {round(p, 3)}"

    # U = {round(u, 3)},

    title = r"r$_{rb}$ = " + f"{round(rbc, n_round)}, p {p}"
    
    raincloud(data, feature1="TC", feature2="ASD",  x_label=label,
              title = title, colors="on", ax=ax[i], size=size_plt)

    i += 1


# %%%                - II.8 - FD & FD balanced             |  ✓  |


rem_everything_bal = {}

size_plt = 15
size_pcp = 13
col_all = "darkslateblue"
col_asd = "lightcoral"
col_tc = "cornflowerblue"


list_fd_max = [0.4]  # , 0.3, 0.2, 0.1
# , "maximum mean FD : 0.3", "maximum mean FD : 0.2", "maximum mean FD : 0.1"
group_labels = ["balanced FD"]

for (this_fd_max, this_label) in zip(list_fd_max, group_labels):

    phenotypic = pd.read_csv(
        'participants/ABIDE_pcp/Phenotypic_V1_0b_preprocessed1.csv')

    # TC outliers:
    funct_measures_tc = graph_measures(
        ts_tc_nr, ts_tc_gsr, prop_threshold=0.15, phenotypic=phenotypic, exclusion=False)

    # Number of participants after non finite exclusion: 454
    # Number of mFC outliers excluded: 22
    # Number of mFD outliers excluded: 6

    # ASD outliers:
    funct_measures_asd = graph_measures(ts_asd_nr_bal, ts_asd_gsr_bal, prop_threshold=0.15, phenotypic=phenotypic, exclusion=False)

    # Number of participants after non finite exclusion: 390
    # Number of mFC outliers excluded: 23
    # Number of mFD outliers excluded: 10

    phenotypic.set_index("SUB_ID", inplace=True)

    # IDs of remaining participants.
    remaining_pcp_id = funct_measures_tc["group_nr"]["ID"].tolist(
    ) + funct_measures_asd["group_nr"]["ID"].tolist()

    # Phenotypic information of remaining participants (ASD + TC).
    rem_pheno = phenotypic.loc[remaining_pcp_id]

    # Remaining TC:
    rem_pheno_tc = rem_pheno.loc[rem_pheno["DX_GROUP"] == 2]

    # Remaining ASD:
    rem_pheno_asd = rem_pheno.loc[rem_pheno["DX_GROUP"] == 1]

    # Extracting the right participants (without separating ASD and TC).
    remaing_pcp_nr = all_pcp_nr.loc[all_pcp_nr["ID"].isin(remaining_pcp_id)]
    remaing_pcp_gsr = all_pcp_gsr.loc[all_pcp_gsr["ID"].isin(remaining_pcp_id)]

    # Put those into a dictionary:
    rem_pcp_nr_gsr = {"nr": remaing_pcp_nr,
                      "gsr": remaing_pcp_gsr}

    # Making sure the index is a list of consecutive numbers.
    rem_pcp_nr_gsr["nr"].index = np.arange(0, len(rem_pcp_nr_gsr["nr"]))
    rem_pcp_nr_gsr["gsr"].index = np.arange(0, len(rem_pcp_nr_gsr["gsr"]))

    # Setting the thresholds:
    thresh = ["0.05", "0.15", "0.25"]

    # Reminder of the parameters required for the analysis:
    # help(stat_analysis)

    id_tc = rem_pheno_tc.index
    id_asd = rem_pheno_asd.index

    # Running the function:

    results = stat_analysis(group=rem_pcp_nr_gsr,
                            id_group1=id_tc,
                            id_group2=id_asd,
                            exclusion=False,
                            # List of thresholds (list of strings).
                            thresholds=thresh,
                            pheno=phenotypic)     # Phenotypic file. Mainly used to exclude
    # mFD outliers.

    # Extracting the graph measures dataframe from the "results" variable.
    meas_asd = results[1]["0.05"]["group_nr"]  # ASD.
    meas_tc = results[0]["0.05"]["group_nr"]  # TC.

    # Concatenate the two to have all the measures of all the participants (both ASD and TC).
    rem_meas_all_pcp = pd.concat([meas_asd, meas_tc])
    # Use the IDs as indices:
    rem_meas_all_pcp = rem_meas_all_pcp.set_index('ID')
    # Order the participants the same way they are ordered in the phenotypic dataframe:
    rem_meas_all_pcp.reindex(index=rem_pheno.index)

    rem_everything = pd.concat([rem_meas_all_pcp, rem_pheno], axis=1)
    rem_everything = rem_everything.sort_values(by="SITE_ID")

    # Make sure the controls are labelled as "0" instead of 2:
    rem_everything['DX_GROUP'][rem_everything['DX_GROUP'] == 2] = 0

    rem_everything["SUB_ID"] = rem_pheno.index

    # Storing this file for every mFD threshold:
    rem_everything_bal[str(0.4)] = rem_everything

    # –––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––
    # –––––––––––––––––– FIGURE 1 : –––––––––––––––––––––––––––––––––––––––––––
    # –––––––––––––––––– FRAMEWISE DISPLACEMENT AND FUNCTIONAL CONNECTIVITY –––
    # –––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––

    # ________ Framewise Displacement Raincloud ________

    # Compare the framewise displacement of both groups
    # just putting the FD in a variable to make it more handy in the mann whitney U
    rem_fd_tc = rem_pheno_tc["func_mean_fd"]
    rem_fd_asd = rem_pheno_asd["func_mean_fd"]
    # len(rem_fd_tc) # checking the the number of participants.
    # len(rem_fd_asd)
    # assessing whether the difference is significant.
    mann_fd = stats.mannwhitneyu(rem_fd_tc, rem_fd_asd)
    # Reporting a few descriptive stats and the p-value.
    print(f"""mean FD TC: {np.median(rem_fd_tc)} and ASD: {np.median(rem_fd_asd)}; 
    median absolute deviation TC: {stats.median_abs_deviation(rem_fd_tc)} and ASD : {stats.median_abs_deviation(rem_fd_asd)} 
    p-value: {mann_fd[1]} """)

    rem_fd = (rem_fd_tc, rem_fd_asd)

    fig, ax = plt.subplots(4, 3, figsize=(14, 14))
    plt.subplots_adjust(top=0.8, bottom=0.01, hspace=0.5, wspace=0.3)
    #plt.tick_params(axis = "both" , labelsize=size_plt)

    # Add a general title to the figure (specifying which mean FD cutoff is used).
    fig.suptitle(this_label,  fontsize=20, fontweight="bold", y=0.86)

    axx = ax[0, 0]

    p_value = mann_fd[1]
    if p_value < 0.01:
        p_value = f"= {Decimal(p_value):.2e}"
    else:
        p_value = f" = {round(p_value, 3)}"
    raincloud(rem_fd, feature1="TC", feature2="ASD",  x_label="Mean FD",
              title=f"p {p_value}", colors="on", ax=axx, size=size_plt)

    # ________ Famewise Displacement Histogram ________

    # FD distribution after balancing : ASD
    ax[1, 0].hist(rem_pheno_asd["func_mean_fd"], bins=80, color=col_asd)
    ax[1, 0].set_xlabel("Mean FD", fontsize=size_plt)
    ax[1, 0].set_ylabel("Frequency", fontsize=size_plt)
    ax[1, 0].text(0.8, 0.9, f'n = {len(rem_pheno_asd["func_mean_fd"])}', horizontalalignment='center',
                  verticalalignment='center', fontsize=size_pcp, transform=ax[1, 0].transAxes)

    # FD distribution after balancing : TC
    ax[2, 0].hist(rem_pheno_tc["func_mean_fd"], bins=80, color=col_tc)
    ax[2, 0].set_xlabel("Mean FD", fontsize=size_plt)
    ax[2, 0].set_ylabel("Frequency", fontsize=size_plt)
    ax[2, 0].text(0.8, 0.9, f'n = {len(rem_pheno_tc["func_mean_fd"])}', horizontalalignment='center',
                  verticalalignment='center', fontsize=size_pcp, transform=ax[2, 0].transAxes)
    ax[1, 0].set_ylim(ax[2, 0].get_ylim())

    # FD distribution after balancing : ALL
    all_mFD = rem_pheno_tc["func_mean_fd"].tolist(
    ) + rem_pheno_asd["func_mean_fd"].tolist()
    ax[3, 0].hist(all_mFD, bins=80, color=col_all)
    ax[3, 0].set_xlabel("Mean FD", fontsize=size_plt)
    ax[3, 0].set_ylabel("Frequency", fontsize=size_plt)
    ax[3, 0].text(0.8, 0.9, f'n = {len(all_mFD)}', horizontalalignment='center',
                  verticalalignment='center', fontsize=size_pcp, transform=ax[3, 0].transAxes)

    # ________ Functional Connectivity Raincloud ________
    axx = ax[0, 1]

    mFC = [results[0]["0.05"]["group_nr"]["mean FC"],
           results[1]["0.05"]["group_nr"]["mean FC"]]
    # makes the titles cleaner on the plot (removes the "Name" and "dtype")
    pval_fc = float(results[3]["p-value"])
    raincloud(mFC, feature1="TC", feature2="ASD", x_label="Mean FC",
              title=f"""p = {round(pval_fc, 3)}""", colors="on", ax=axx, size=size_plt)

    # ________ Functional Connectivity Histogram ________

    #ax.get_shared_y_axes().join(ax[1,1], ax[2,1], ax[3,1])

    # FC distribution : ASD
    ax[1, 1].hist(results[1]["0.05"]["group_nr"]
                  ["mean FC"], bins=80, color=col_asd)
    ax[1, 1].set_xlabel("Mean FC", fontsize=size_plt)
    ax[1, 1].set_ylabel("Frequency", fontsize=size_plt)
    ax[1, 1].text(0.8, 0.9, f'n = {len(results[1]["0.05"]["group_nr"]["mean FC"])}', horizontalalignment='center',
                  verticalalignment='center', fontsize=size_pcp, transform=ax[1, 1].transAxes)

    # FC distribution : TC
    ax[2, 1].hist(results[0]["0.05"]["group_nr"]
                  ["mean FC"], bins=80, color=col_tc)
    ax[2, 1].set_xlabel("Mean FC", fontsize=size_plt)
    ax[2, 1].set_ylabel("Frequency", fontsize=size_plt)
    ax[2, 1].text(0.8, 0.9, f'n = {len(results[0]["0.05"]["group_nr"]["mean FC"])}', horizontalalignment='center',
                  verticalalignment='center', fontsize=size_pcp, transform=ax[2, 1].transAxes)
    ax[1, 1].set_ylim(ax[2, 1].get_ylim())

    # FC distribution : ALL
    # convert into list because it was in pandas core series so it wasn't working.
    all_mFC = results[0]["0.05"]["group_nr"]["mean FC"].tolist(
    ) + results[1]["0.05"]["group_nr"]["mean FC"].tolist()
    ax[3, 1].hist(all_mFC, bins=80, color=col_all)
    ax[3, 1].set_xlabel("Mean FC", fontsize=size_plt)
    ax[3, 1].set_ylabel("Frequency", fontsize=size_plt)
    ax[3, 1].text(0.8, 0.9, f'n = {len(all_mFC)}', horizontalalignment='center',
                  verticalalignment='center', fontsize=size_pcp, transform=ax[3, 1].transAxes)

    # ________ Correlation FD and FC ________
    # ASD
    sns.regplot(x=rem_pheno_asd["func_mean_fd"], y=mFC[1], scatter_kws={
                's': 8}, ax=ax[1, 2], color=col_asd)
    ax[1, 2].set_xlabel("Mean FD", fontsize=size_plt)
    ax[1, 2].set_ylabel("Mean FC", fontsize=size_plt)
    ax[1, 2].text(0.8, 0.9, f'n = {len(mFC[1])}', horizontalalignment='center',
                  fontsize=size_pcp, verticalalignment='center', transform=ax[1, 2].transAxes)
    r, p = stats.pearsonr(x=rem_pheno_asd["func_mean_fd"], y=mFC[1])
    # round them cause we don't need a tone of decimals
    ax[1, 2].set_title(
        f"r = {round(r, 3)}, p = {round(p, 3)}", fontsize=size_plt)

    if p < 0.01:
        p = f"= {Decimal(p):.2e}"
    else:
        p = f" = {round(p, 3)}"

    ax[1, 2].set_title(f"r = {round(r, 3)}, p {p}", fontsize=size_plt)

    # TC
    sns.regplot(x=rem_pheno_tc["func_mean_fd"], y=mFC[0], scatter_kws={
                's': 8}, ax=ax[2, 2], color=col_tc)
    ax[2, 2].set_xlabel("Mean FD", fontsize=size_plt)
    ax[2, 2].set_ylabel("Mean FC", fontsize=size_plt)
    ax[2, 2].text(0.8, 0.9, f'n = {len(mFC[0])}', horizontalalignment='center',
                  verticalalignment='center', fontsize=size_pcp, transform=ax[2, 2].transAxes)
    r, p = stats.pearsonr(x=rem_pheno_tc["func_mean_fd"], y=mFC[0])

    if p < 0.01:
        p = f"= {Decimal(p):.2e}"
    else:
        p = f" = {round(p, 3)}"

    ax[2, 2].set_title(f"r = {round(r, 3)}, p {p}", fontsize=size_plt)

    # ALL
    sns.regplot(x=all_mFD, y=all_mFC, scatter_kws={
                's': 8}, ax=ax[3, 2], color=col_all)
    ax[3, 2].set_xlabel("Mean FD", fontsize=size_plt)
    ax[3, 2].set_ylabel("Mean FC", fontsize=size_plt)
    ax[3, 2].text(0.8, 0.9, f'n = {len(all_mFC)}', horizontalalignment='center',
                  verticalalignment='center', fontsize=size_pcp, transform=ax[3, 2].transAxes)
    r, p = stats.pearsonr(x=all_mFD, y=all_mFC)

    if p < 0.01:
        p = f"= {Decimal(p):.2e}"
    else:
        p = f" = {round(p, 3)}"

    ax[3, 2].set_title(f"r = {round(r, 3)}, p {p}", fontsize=size_plt)

    # ________ Correlation FD and FC as a function of distance ________

    # FD from all remaining participants.
    rem_pcp_fd = rem_everything["func_mean_fd"]

    # Vectorised upper triangular of remaining participants:
    rem_vec_all = rem_everything["vectorised matrix"]

    rem_vec_all = np.stack(rem_vec_all, axis=0)

    # Calculate the correlation between edges and FD accross participants.
    corr_fd_fc = np.zeros(rem_vec_all.shape[1])
    for i in range(rem_vec_all.shape[1]):
        corr_fd_fc[i] = stats.pearsonr(rem_vec_all[:, i], rem_pcp_fd)[0]

    # Put the correlation coefficients and the distances in one dataframe.
    corr_dist = pd.DataFrame(
        {"corr_fd_fc": corr_fd_fc, "distance": cond_dist_mat})
    corr_dist = corr_dist.dropna()  # Get rid of the rows that have NaN values.

    values = np.vstack([corr_dist["corr_fd_fc"], corr_dist["distance"]])
    kernel = stats.gaussian_kde(values)(values)

    sns.scatterplot(corr_dist,
                    x="distance",
                    y="corr_fd_fc",
                    c=kernel,
                    cmap="viridis",
                    ax=ax[0, 2])

    # NOT SURE ABOUT FONTSIZE HERE
    sns.regplot(data=corr_dist, x="distance", y="corr_fd_fc",
                scatter=False, ax=ax[0, 2], color="red")
    ax[0, 2].set_xlabel("Distance", fontsize=size_plt)
    ax[0, 2].set_ylabel("Corr mFD and mFC", fontsize=size_plt)
    r, p = stats.pearsonr(x=corr_dist["distance"], y=corr_dist["corr_fd_fc"])

    if p < 0.01:
        p = "< 0.01"
    else:
        p = f" = {round(p, 3)}"

    # round them cause we don't need a tone of decimals
    ax[0, 2].set_title(f"r = {round(r, 5)}, p {p}", fontsize=size_plt)

    # –––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––
    # –––––––––––––––––– FIGURE 2 : –––––––––––––––––––––––––––––––––––––––––––
    # –––––––––––––––––– EDGES WEIGHTS FOR NR, MR and GSR –––––––––––––––––––––
    # –––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––

    # Average the vectorised upper triangulars of the connectivity matrices
    # for all three conditions (NR, MR and GSR).
    vec_nr = results[0]["0.05"]["group_nr"]["vectorised matrix"].tolist(
    ) + results[1]["0.05"]["group_nr"]["vectorised matrix"].tolist()
    vec_gsr = results[0]["0.05"]["group_gsr"]["vectorised matrix"].tolist(
    ) + results[1]["0.05"]["group_gsr"]["vectorised matrix"].tolist()
    vec_mreg = results[0]["0.05"]["group_mreg"]["vectorised matrix"].tolist(
    ) + results[1]["0.05"]["group_mreg"]["vectorised matrix"].tolist()

    # print(f"THIS IS IMPORTANTS !!! vectors: {len(vec_nr)} should be the num of all pcp and {len(vec_nr[0])} should be the num of edges and the type is {type(vec_nr)}")

    mfc_nr = np.mean(vec_nr, axis=0)
    mfc_gsr = np.mean(vec_gsr, axis=0)
    mfc_mreg = np.mean(vec_mreg, axis=0)

    # print(f"THIS IS IMPORTANTS !!! averaged vectors : {len(mfc_nr)} should be the num of edges and the type is {type(mfc_nr)}")

    # Create a figure with six subplots.
    fig2, ax2 = plt.subplots(2, 3, figsize=(14, 10))
    fig2.subplots_adjust(top=0.6, bottom=0.01, hspace=0.5, wspace=0.3)
    # Add a general title to the figure.
    fig2.suptitle(this_label,  fontsize=20, fontweight="bold", y=0.68)
    ax2[0, 0].text(x=0.1, y=1,  s="A",  fontsize=20, fontweight="bold")
    ax2[1, 0].text(x=0.1, y=355, s="B",  fontsize=20, fontweight="bold")

    # GSR in function of NR:
    ax2[0, 0].scatter(x=mfc_nr, y=mfc_gsr, s=2, c=col_all)
    ax2[0, 0].set_xlabel("NR", fontsize=14)
    ax2[0, 0].set_ylabel("GSR", fontsize=14)
    ax2[0, 0].text(0.15, 0.9, f'n = {len(vec_nr)}', horizontalalignment='center',
                   verticalalignment='center', fontsize=size_pcp, transform=ax2[0, 0].transAxes)

    # MR in function of NR:
    ax2[0, 1].scatter(x=mfc_nr, y=mfc_mreg, s=2, c=col_all)
    ax2[0, 1].set_xlabel("NR", fontsize=14)  # averaged weights
    ax2[0, 1].set_ylabel("MR", fontsize=14)  # averaged weights
    ax2[0, 1].text(0.15, 0.9, f'n = {len(vec_mreg)}', horizontalalignment='center',
                   verticalalignment='center', fontsize=size_pcp, transform=ax2[0, 1].transAxes)

    # MR in function of GSR:
    ax2[0, 2].scatter(x=mfc_gsr, y=mfc_mreg, s=2, c=col_all)
    ax2[0, 2].set_xlabel("GSR", fontsize=14)
    ax2[0, 2].set_ylabel("MR", fontsize=14)
    ax2[0, 2].text(0.15, 0.9, f'n = {len(vec_gsr)}', horizontalalignment='center',
                   verticalalignment='center', fontsize=size_pcp, transform=ax2[0, 2].transAxes)

    # ____ edge distribution ____

    # print(f"THIS IS ALSO IMPORTANT (flattened): {len(flat_mfc_nr)} and type {type(flat_mfc_nr)}")

    ax2[1, 0].hist(mfc_nr, bins=70, color=col_all)
    ax2[1, 0].set_xlabel("NR", fontsize=14)
    # Add the sample size on the graph
    ax2[1, 0].text(0.15, 0.9, f'n = {len(vec_nr)}', horizontalalignment='center',
                   verticalalignment='center', fontsize=size_pcp, transform=ax2[1, 0].transAxes)

    ax2[1, 1].hist(mfc_mreg, bins=70, color=col_all)
    ax2[1, 1].set_xlabel("MR", fontsize=14)
    ax2[1, 1].text(0.15, 0.9, f'n = {len(vec_mreg)}', horizontalalignment='center',
                   verticalalignment='center', fontsize=size_pcp, transform=ax2[1, 1].transAxes)

    ax2[1, 2].hist(mfc_gsr, bins=70, color=col_all)
    ax2[1, 2].set_xlabel("GSR", fontsize=14)
    ax2[1, 2].text(0.15, 0.9, f'n = {len(vec_gsr)}', horizontalalignment='center',
                   verticalalignment='center', fontsize=size_pcp, transform=ax2[1, 2].transAxes)
    ax2[1, 0].set_ylim(ax2[1, 2].get_ylim())
    ax2[1, 1].set_ylim(ax2[1, 2].get_ylim())


# %%           - III - CLASSIFICATION                      –––––––
# %%%                - III.1 - Scanning sites              |  ✓  | to clean up

cmap = cmap = plt.get_cmap('coolwarm')
plt.set_cmap(cmap)

rem_everything_un = rem_everything_un2
# Ordering the data frame by scanning site (very likely has been done higher up but just in case):
rem_everything_un = rem_everything_un.sort_values(by="SITE_ID")


# Loop through all the unique values of SITE_ID (so through every possible scanning
# site), and sum them to have the numer of participants corresponding to each site.
sum_sites = [sum(rem_everything_un["SITE_ID"] == i)
             for i in np.unique(rem_everything_un["SITE_ID"])]


# Visualise the correlation of the vectorise upper triangular of the connectivity matrix
# between each pair of participant.
fc_vectors = rem_everything_un["vectorised matrix"].tolist()
corr_mat = np.corrcoef(fc_vectors)
plt.imshow(corr_mat)
plt.show()


# Trying the seaborn heatmap method be the result is a bit weird:
sns.heatmap(pd.DataFrame(corr_mat))


# Modified version of FV's function to be able to add the lines to delimitate
# the scanning sites directly on the imshow plot.
def add_subnetwork_lines(hm, roi_nums):
    corr_mat = np.corrcoef(hm)
    plt.imshow(corr_mat)
    plt.colorbar()

    for i in np.cumsum(roi_nums):
        # print(i)
        plt.axvline(i - 0.25, color="black", linestyle="-", linewidth=0.2)
        plt.axhline(i - 0.25, color="black", linestyle="-", linewidth=0.2)

    # Define the limits of the axes, otherwise the lines are "too long".
    plt.xlim(0, len(corr_mat[0]))
    # Here I do it the other way round so that the
    plt.ylim(len(corr_mat[1]), 0)
    # The axis stays ordered the same way and is not reverse.
    plt.xticks([])
    plt.yticks([])
    
    plt.show()


add_subnetwork_lines(fc_vectors, sum_sites)  # Plotting the matrix.

# POOL SOME SITE TOGETHER

# - Now that the visualisation is done, let's pool the scanning site with the
# least participants with the scanning site that has the most similar average
# correlation.
where_min = np.where(sum_sites == np.min(sum_sites))
min_site = np.unique(rem_everything_un["SITE_ID"])[where_min]
print(f"the scanning site that has the less participants is {min_site}")

# Splitting the correlation matrix to get chunks of correlation coefficients
# by scanning site:
first_split = np.split(corr_mat, np.cumsum(sum_sites), axis=0)

# Splitting those chunk once again by scanning site to have a chunk for every
# combination of scanning site (basically all the squares we see on the plot):
second_split = [np.split(this_split, np.cumsum(sum_sites), axis=1)
                for this_split in first_split]


# Averaging each squares:
mn_corr_site = [np.mean(this_array)
                for this_list in second_split for this_array in this_list]
# Get rid of nan values (how come I get nans????????????):
mn_corr_site = [this_mn for this_mn in mn_corr_site if str(this_mn) != "nan"]
# Reshaping into a site x site matrix:
mn_corr_site = np.array(mn_corr_site).reshape(len(sum_sites), len(sum_sites))
# Potting the result:
plt.imshow(mn_corr_site)
plt_size = 10
plt.xticks(range(len(sum_sites)), np.unique(rem_everything_un["SITE_ID"]).tolist(), rotation=90, fontsize = plt_size)
plt.yticks(range(len(sum_sites)), np.unique(rem_everything_un["SITE_ID"]).tolist() , fontsize = plt_size)
plt.colorbar()

# mn_corr_site[1,:] = 1 # checking that i'm indexing the right site (CMU) -> yes :)
# plt.imshow(mn_corr_site)

# get the second largest average (the first largest being the CMU-CMU):
sort_cmu = sorted(mn_corr_site[1, :])[-2]
# See which site corresponds to this value:
# keep only the integer (it was a tuple -> array -> integer)
max_site = np.where(mn_corr_site[1, :] == sort_cmu)[0][0]
print("the site with the strongest average correlation is:", np.unique(
    rem_everything_un["SITE_ID"]).tolist()[max_site])  # indexing to see which site it corresponds to.


rem_everything_un["SITE_ID"][rem_everything_un["SITE_ID"] == "USM"] = "USM-CMU"
rem_everything_un["SITE_ID"][rem_everything_un["SITE_ID"] == "CMU"] = "USM-CMU"

rem_everything_un["stratify"] = rem_everything_un['DX_GROUP'].astype(
    str) + "-" + rem_everything_un["SITE_ID"]


# %%%                - III.1 - Stratify & balance          |  ✓  |


the_group = rem_everything_un2

print(f"""Number of controls in the unbalanced group: {sum(the_group["DX_GROUP"] == 0)}.
Number of ASD patients in the unbalanced group: {sum(the_group["DX_GROUP"] == 1)}. 
Number of ASD participants to add: {sum(the_group["DX_GROUP"] == 0) - sum(the_group["DX_GROUP"] == 1)}""")

asd_rows = the_group[the_group["DX_GROUP"] == 1][:70]
rem_everyt_equal = pd.concat([the_group, asd_rows], axis=0)
rem_everyt_equal = rem_everyt_equal.reset_index()

print(f"""Number of controls in the unbalanced group: {sum(rem_everyt_equal["DX_GROUP"] == 0)}.
Number of ASD patients in the unbalanced group: {sum(rem_everyt_equal["DX_GROUP"] == 1)}. 
Number of ASD participants to add: {sum(rem_everyt_equal["DX_GROUP"] == 0) - sum(rem_everyt_equal["DX_GROUP"] == 1)}""")


# # Put the IDs back to being a column instead of the index:
# rem_everyt_equal["SUB_ID"] = rem_everyt_equal.index
# # Make sure the index is made of consecutive numbers:
# rem_everyt_equal = rem_everyt_equal.reset_index()

X = rem_everyt_equal
rem_everyt_equal["stratify"] = X["DX_GROUP"].astype(str) + X["SITE_ID"].astype(str)


# Divide train and test data:

df_trainval, df_test = train_test_split(rem_everyt_equal,  # Data being split.
                                        test_size=0.2,   # 80-20% separation.
                                        # Column use for statification:
                                        # (diagnosis-scanning site combination)
                                        stratify=rem_everyt_equal[["stratify"]])


# Divide the train data into 5 train/validation folds:
X = df_trainval             # Data to split.
# Stratification target (diagnosis-site combination).
y = df_trainval["stratify"]
nfolds = 5                  # Number of cross-validation folds.
skf = StratifiedKFold(n_splits=nfolds)
skf.get_n_splits(X, y)      # Check that the right number of folds was set.


train = []      # Will store the training data.
validation = []  # Will store the validation data.

StratifiedKFold(n_splits=nfolds, random_state=None, shuffle=False)

for train_index, validation_index in skf.split(X, y):

    train += [X.iloc[train_index]]
    validation += [X.iloc[validation_index]]


# Checking that there is the same number of patients and sites in the groups:
# It's not exactly the same is it normal ????????????????????????????????????
print(len(np.where(train[2]["DX_GROUP"] == 1)[0]))
print(len(np.where(train[2]["DX_GROUP"] == 0)[0]))
print(len(np.where(train[0]["stratify"] == "0USM-CMU")[0]))
print(len(np.where(train[4]["stratify"] == "0USM-CMU")[0]))


# %%%                - III.1 - training              |  ✓  |


phenotypic = pd.read_csv('participants/ABIDE_pcp/Phenotypic_V1_0b_preprocessed1.csv')

# Initialise a nested default dictionary to store the accuracy of the different
# function, thresholds, and folds.
accuracy = defaultdict(dict)
specificity = defaultdict(dict)
sensitivity = defaultdict(dict)
recall = defaultdict(dict)
precision = defaultdict(dict)


scaler = StandardScaler()
kernels = ["rbf", "poly", "linear", "sigmoid"]
thresholds = [0.05, 0.15, 0.25]
groups = ["group_nr", "group_gsr", "group_mreg"]
c_params = np.arange(0.1, 2.1, 0.1)
gam_params = np.arange(0.1, 1.1, 0.05)


for i in range(nfolds):  # Looping through the 5 folds.
        
    # –––––––––––––––––––– DATA PREPARATION ––––––––––––––––––––

    # Get the time series of the participants in the first fold:
    df_train_nr = all_pcp_nr[all_pcp_nr["ID"].isin(train[i]["subject"])]
    # Use the participants' ID as index:
    df_train_nr = df_train_nr.set_index('ID')
    # Reorder the time series so that they match the order of the training data:
    df_train_nr = df_train_nr.reindex(index=train[i]['subject'])
    # Resetting the index so that it's back to being consecutive numbers
    # (0, 1, ..., N):
    df_train_nr = df_train_nr.reset_index()
    # "reindex" part switch the ID columns to being "SUB_ID". It needs to be
    # switched back to "ID" for the "graph_measures" function to work:
    df_train_nr = df_train_nr.rename(columns={'subject': "ID"})
    
    
    # Doing the exact same thing, but for the GSR time series:
    df_train_gsr = all_pcp_gsr[all_pcp_gsr["ID"].isin(train[i]["subject"])]
    df_train_gsr = df_train_gsr.set_index('ID')
    df_train_gsr = df_train_gsr.reindex(index=train[i]['subject'])
    df_train_gsr = df_train_gsr.reset_index()
    df_train_gsr = df_train_gsr.rename(columns={'subject': "ID"})
    
    # Once again doing the same thing, but for the NR validation fold...
    df_val_nr = all_pcp_nr[all_pcp_nr["ID"].isin(validation[i]["subject"])]
    df_val_nr = df_val_nr.set_index('ID')
    df_val_nr = df_val_nr.reindex(index=validation[i]['subject'])
    df_val_nr = df_val_nr.reset_index()
    df_val_nr = df_val_nr.rename(columns={'subject': "ID"})

    # ...and for the GSR validation fold.
    df_val_gsr = all_pcp_gsr[all_pcp_gsr["ID"].isin(
        validation[i]["subject"])]
    df_val_gsr = df_val_gsr.set_index('ID')
    df_val_gsr = df_val_gsr.reindex(index=validation[i]['subject'])
    df_val_gsr = df_val_gsr.reset_index()
    df_val_gsr = df_val_gsr.rename(columns={'subject': "ID"})

   # –––––––––––––––––––– CLASSIFICATION ––––––––––––––––––––

    # Looping through the thresholds to be able to run the classification for
    # each of them:
    for this_threshold in thresholds:

        train_folds_nr_gsr_mreg = graph_measures(df_train_nr, df_train_gsr, prop_threshold = this_threshold, phenotypic = phenotypic, exclusion = False)
        
        
        print("___________________ ONE THRESHOLD ___________________")
        
        # Looping throught the pipelines:
        for this_group in groups:
            
            x = train_folds_nr_gsr_mreg[this_group] #[["cluster", "assort", "charpath"]]
            x["vect_matrix"] = x["vectorised matrix"]
            x = pd.DataFrame(x["vect_matrix"])
            x[np.arange(len(x["vect_matrix"][0]))] = pd.DataFrame(x.vect_matrix.tolist(), index= x.index)
            x = x.drop('vect_matrix', axis=1)
            

            if this_group == "group_mreg":
                val_folds_gsr_pred = graph_measures(df_train_nr,
                                                    df_val_gsr,
                                                    prop_threshold=this_threshold,
                                                    phenotypic=phenotypic,
                                                    exclusion=False,
                                                    group_pred=df_val_nr)

                val_fold = val_folds_gsr_pred["group_pred"] #[["cluster", "assort", "charpath"]]

            elif this_group == "group_gsr":

                val_folds_gsr_mreg = graph_measures(df_val_nr,
                                                    df_val_gsr,
                                                    prop_threshold=this_threshold,
                                                    phenotypic=phenotypic,
                                                    exclusion=False)

                val_fold = val_folds_gsr_mreg["group_gsr"] #[["cluster", "assort", "charpath"]]

            elif this_group == "group_nr":
                val_folds_nr = graph_measures(df_val_nr,
                                              df_val_gsr,
                                              prop_threshold=this_threshold,
                                              phenotypic=phenotypic,
                                              exclusion=False)

                val_fold = val_folds_nr["group_nr"] #[["cluster", "assort", "charpath"]]
                
            # val_fold = val_fold[["cluster", "assort", "charpath"]]
            val_fold["vect_matrix"] = val_fold["vectorised matrix"]
            val_fold = pd.DataFrame(val_fold["vect_matrix"])
            val_fold[np.arange(len(val_fold["vect_matrix"][0]))] = pd.DataFrame(val_fold.vect_matrix.tolist(), index= val_fold.index)
            val_fold = val_fold.drop('vect_matrix', axis=1)
            
            print("___________________ ONE GROUP ___________________")
            
            # –––––––––––––––––––– RESULTS ––––––––––––––––––––
            
            for this_kernel in kernels:

                for this_c in c_params: 
                    
                    for this_gam in gam_params:
                        
                        svc = SVC(kernel=this_kernel, C = this_c, gamma = this_gam, probability=True)  # Setting the rbf SVM.
                        
                        # Combining the scaler and the rbf SVM into a pipeline:
                        pipe_svm = make_pipeline(scaler, svc)
                       
                        pipe_svm.fit(x, train[i]["DX_GROUP"])
                        
                        # Prediction using the validation fold:
                        val_pred = pipe_svm.predict(val_fold)
        
                        combination = (this_kernel + "_" + "C" + str(this_c) + "_" + "Gam" + str(this_gam) + "_" + this_group + "_" + str(this_threshold) )
        
                        # Storing the accuracy score for every pipeline:
                        accuracy[combination][i] = accuracy_score(validation[i]["DX_GROUP"], val_pred)
                        # Recall score:
                        recall[combination][i] = recall_score(validation[i]["DX_GROUP"], val_pred)
                        # Precision score:
                        precision[combination][i] = precision_score(validation[i]["DX_GROUP"], val_pred)
        
        
                        # Sentitivity and specificity - for each pipeline:
                        # Not sure it's the right way to do???????????????
                        tn, fp, fn, tp = confusion_matrix(validation[i]["DX_GROUP"], val_pred).ravel()
                        specificity[combination][i] = tn / (tn + fp)
                        sensitivity[combination][i] = tp / (tp + fn)
                        
                        print("___________________ ONE PARAMETER ___________________")

    print("___________________ ! ONE FOLD ! ___________________")
        


# create a new var to store the results of the vectorised matrix + store the other results !!!!!!!!!!!!!!!!!

columns = [accuracy , recall , precision , specificity , sensitivity]
columns2 = ["accuracy" , "recall" , "precision" , "specificity" , "sensitivity"]

df_all_results = pd.DataFrame()
for i in range(len(columns)):                 # Loop through the 7 scores. 
    scores_df = pd.DataFrame(columns[i])      # Convert defaultdict into df. 
    scores_df = scores_df.transpose()         # (folds, pipeline) -> (pipeline, fold).
    scores_df = np.mean(scores_df, axis = 1)  # Average the 5 folds.
    df_all_results[columns2[i]] = scores_df   # Add the averages in a dataframe.

df_all_results_saved = df_all_results

# def explode_array_column(row):
#     return pd.Series(row['vectorised matrix'])

# expanded_cols = x.apply(explode_array_column, axis=1)
# expanded_cols.columns = ['col_{}'.format(i) for i in range(expanded_cols.shape[1])]

# rem_everyt_equal["vect_matrix"] = rem_everyt_equal["vectorised matrix"]

# trius = pd.DataFrame()
# trius[np.arange(len(rem_everyt_equal["vect_matrix"][300]))] = pd.DataFrame(rem_everyt_equal.vect_matrix.tolist(), index= rem_everyt_equal.index)
# trius.head()

# df = pd.concat([x, expanded_cols], axis=1)
# df = df.drop('vectorised matrix', axis=1)

# print(df)


df_all_results = df_all_results.sort_values('accuracy', ascending=False)

pipelines = pd.DataFrame(df_all_results.index.str.split('_').tolist(), columns = np.arange(6))

df_all_results = df_all_results.reset_index()

df_all_results = df_all_results.join(pipelines)

df_all_results["min_pip"] = df_all_results[0].astype(str) + "_" + df_all_results[4].astype(str) + "_" + df_all_results[5].astype(str)

max_grid = pd.DataFrame()

#[max_grid[i] = df_all_results.loc[df_all_results.loc[df_all_results["min_pip"] == i]]] #.accuracy.idxmax()] for i in df_all_results["min_pip"]] 

min_pipelines = np.unique(df_all_results["min_pip"])
for this_pip in min_pipelines:
    
    max_grid[this_pip] = df_all_results.loc[df_all_results.loc[df_all_results["min_pip"] == this_pip].accuracy.idxmax()]

max_grid = max_grid.transpose()

max_grid = max_grid.sort_values('min_pip', ascending=False)
max_grid = max_grid.sort_values('accuracy', ascending=False)

max_grid2_v1 = pd.DataFrame()
max_grid2_v1[["kernel", "group", "threshold", "accuracy", "precision", "sensitivity", "specificity"]] = max_grid[[0, 4, 5, "accuracy", "precision", "sensitivity", "specificity"]]
max_grid2_v1[["accuracy", "precision", "sensitivity", "specificity"]] = max_grid2_v1[["accuracy", "precision", "sensitivity", "specificity"]].astype(float).round(3)

#df_all_results = df_all_results.sort_values('accuracy', ascending=False)

max_grid_sortacc = max_grid.sort_values('accuracy', ascending=False)

max_grid_thres = pd.DataFrame()
for kernel in np.unique(max_grid2_v1["kernel"]):
    for group in np.unique(max_grid2_v1["group"]):
        combination = kernel + "_" + group
        max_grid_thres[combination] = max_grid[(max_grid2_v1["kernel"] == kernel) & (max_grid2_v1["group"] == group)][["accuracy", "precision", "sensitivity", "specificity"]].mean(axis=0)


max_grid_thres = max_grid_thres.transpose()
max_thres_pip = pd.DataFrame(max_grid_thres.index.str.split('_').tolist(), columns = np.arange(2))
max_grid_thres = max_grid_thres.reset_index()
max_grid_thres = max_grid_thres.join(max_thres_pip)
max_grid_thres2 = pd.DataFrame()
max_grid_thres2[["kernel", "group", "accuracy", "precision", "sensitivity", "specificity"]] = max_grid_thres[[0, 1, "accuracy", "precision", "sensitivity", "specificity"]]
max_grid_thres2 = max_grid_thres2.sort_values(["kernel", "group"], ascending=False)
max_grid_thres2[["accuracy", "precision", "sensitivity", "specificity"]] = max_grid_thres2[["accuracy", "precision", "sensitivity", "specificity"]].astype(float).round(3)
max_grid_thres2 = max_grid_thres2.sort_values("accuracy", ascending=False)

# #____________

# with open(r"train_un_final.pickle", "wb") as output_file:
#     cPickle.dump(df_all_results, output_file)


# df_all_results = df_all_results.round(3)
# df_all_results["pipeline"] = df_all_results.index

max_grid = max_grid.reset_index()

import docx
saved = max_grid
# Initialise the Word document
doc = docx.Document()
# Initialise the table
t = doc.add_table(rows=1, cols=saved.shape[1])
# Add borders
t.style = 'TableGrid'
# Add the column headings
for j in range(saved.shape[1]):
    t.cell(0, j).text = saved.columns[j]
# Add the body of the data frame
for i in range(saved.shape[0]):
    row = t.add_row()
    for j in range(saved.shape[1]):
        cell = saved.iat[i, j]
        row.cells[j].text = str(cell)
# Save the Word doc
doc.save('table_crossval_meas_un_hyperparam.docx')



max_grid_accuracy = max_grid.sort_values('accuracy', ascending=False)

# FILES : 
    # mistake : un -> bal 
    
    # BALANCED
# table_crossval_meas_un_hyperparam
# table_crossval_meas_un_allclean
# table_crossval_meas_un_thres

    # UPPER TRIANGULAR BALANCED
# table_crossval_meas_unbal_hyperparam
# table_crossval_meas_unbal_allclean
# table_crossval_meas_unbal_thres

    # UNBALANCED
# table_crossval_meas_unbal_hyperparam
# table_crossval_meas_unbal_allclean
# table_crossval_meas_unbal_thres




#%% TESTING

auc = {}
from matplotlib import cm
from sklearn.metrics import confusion_matrix, ConfusionMatrixDisplay



phenotypic = pd.read_csv('participants/ABIDE_pcp/Phenotypic_V1_0b_preprocessed1.csv')



# Get the time series of the participants in the first fold:
df_train_nr = all_pcp_nr[all_pcp_nr["ID"].isin(df_trainval["subject"])]
# Use the participants' ID as index:
df_train_nr = df_train_nr.set_index('ID')
# Reorder the time series so that they match the order of the training data:
df_train_nr = df_train_nr.reindex(index=df_trainval['subject'])
# Resetting the index so that it's back to being consecutive numbers
# (0, 1, ..., N):
df_train_nr = df_train_nr.reset_index()
# "reindex" part switch the ID columns to being "SUB_ID". It needs to be
# switched back to "ID" for the "graph_measures" function to work:
df_train_nr = df_train_nr.rename(columns={'subject': "ID"})

# Doing the exact same thing, but for the GSR time series:
df_train_gsr = all_pcp_gsr[all_pcp_gsr["ID"].isin(df_trainval["subject"])]
df_train_gsr = df_train_gsr.set_index('ID')
df_train_gsr = df_train_gsr.reindex(index=df_trainval['subject'])
df_train_gsr = df_train_gsr.reset_index()
df_train_gsr = df_train_gsr.rename(columns={'subject': "ID"})

# Once again doing the same thing, but for the NR validation fold...
df_val_nr = all_pcp_nr[all_pcp_nr["ID"].isin(df_test["subject"])]
df_val_nr = df_val_nr.set_index('ID')
df_val_nr = df_val_nr.reindex(index=df_test['subject'])
df_val_nr = df_val_nr.reset_index()
df_val_nr = df_val_nr.rename(columns={'subject': "ID"})

# ...and for the GSR validation fold.
df_val_gsr = all_pcp_gsr[all_pcp_gsr["ID"].isin(df_test["subject"])]
df_val_gsr = df_val_gsr.set_index('ID')
df_val_gsr = df_val_gsr.reindex(index=df_test['subject'])
df_val_gsr = df_val_gsr.reset_index()
df_val_gsr = df_val_gsr.rename(columns={'subject': "ID"})



 
# # Fill in the pipeline of interest here: 
# scaler = StandardScaler()
# this_kernel = "rbf"
# this_threshold = 0.25 # 0.25] # 0.05, , 0.25
# this_group = "group_nr"
# group_label = "NR"
# this_c = 0.7
# this_gam = 0.9

# scaler = StandardScaler()
# this_kernel = "sigmoid"
# this_threshold = 0.25 # 0.25] # 0.05, , 0.25
# this_group = "group_mreg"
# group_label = "NR"
# this_c = 1
# this_gam = 1


scaler = StandardScaler()
this_kernel = "rbf"
this_threshold = 0.15 # 0.25] # 0.05, , 0.25
this_group = "group_nr"
group_label = "NR"
this_c = 1
this_gam = 1


# –––––––––––––––––––– CLASSIFICATION ––––––––––––––––––––

# Looping through the thresholds to be able to run the classification for
# each of them:


train_folds_nr_gsr_mreg = graph_measures(df_train_nr, df_train_gsr, prop_threshold = this_threshold, phenotypic = phenotypic, exclusion = False)


#–––––––––

val_folds_gsr_pred = graph_measures(df_train_nr,
df_val_gsr,
prop_threshold=this_threshold,
phenotypic=phenotypic,
exclusion=False,
group_pred=df_val_nr)
   
val_fold = val_folds_gsr_pred["group_pred"][["cluster", "assort", "charpath"]]   
x = train_folds_nr_gsr_mreg[this_group][["cluster", "assort", "charpath"]]


# Combining the scaler and the rbf SVM into a pipeline:
svc = SVC(kernel=this_kernel, C = this_c, gamma = this_gam, probability=True)  # Setting the rbf SVM.
pipe_svm = make_pipeline(scaler, svc)
# Fitting the pipeline
pipe_svm.fit(x, df_trainval["DX_GROUP"])

# Prediction using the validation fold:
val_pred = pipe_svm.predict(val_fold)
val_proba = pipe_svm.predict_proba(val_fold)

combination = (this_kernel + "_" + "C" + str(this_c) + "_" + "Gam" + str(this_gam) + "_" + this_group + "_" + str(this_threshold) )

# Storing the accuracy score for every pipeline:
accuracy = accuracy_score(df_test["DX_GROUP"], val_pred)
# Recall score:
recall = recall_score(df_test["DX_GROUP"], val_pred)
# Precision score:
precision = precision_score(df_test["DX_GROUP"], val_pred)


# Sentitivity and specificity - for each pipeline:
# Not sure it's the right way to do???????????????
tn, fp, fn, tp = confusion_matrix(df_test["DX_GROUP"], val_pred).ravel()
specificity = tn / (tn + fp)
sensitivity = tp / (tp + fn)

# Code from week 3:

cm = confusion_matrix(df_test["DX_GROUP"], val_pred)
# Plot
disp = ConfusionMatrixDisplay(confusion_matrix=cm, display_labels=pipe_svm.classes_)
disp.plot();
plt.title(f"{this_kernel}; {group_label}, {this_threshold}%")
plt.show()
   
fpr, tpr, thresholds = roc_curve(df_test["DX_GROUP"], val_proba[:,1])
auc = roc_auc_score(df_test["DX_GROUP"], val_proba[:,1])
plt.plot(fpr, tpr, linestyle='--');
plt.title(f'{this_kernel} kSVM; {group_label}, {this_threshold}%; AUC = {round(auc,4)}');
plt.xlabel('False Positive Rate = 1 - Specificity');
plt.ylabel('True Positive Rate = Sensitivity');
plt.show()
   

# –––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––

df_all_results_test = pd.DataFrame({"combination": combination,
                                    "accuracy": accuracy,
                                    "precision": precision,
                                    "sensitivity": sensitivity,
                                    "specificity": specificity}, 
                                    index=[0])


df_all_results_test3 = df_all_results_test.round(3).transpose()

# –––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––

# Save output

import docx
saved = df_all_results_test1
# Initialise the Word document
doc = docx.Document()
# Initialise the table
t = doc.add_table(rows=1, cols=saved.shape[1])
# Add borders
t.style = 'TableGrid'
# Add the column headings
for j in range(saved.shape[1]):
    t.cell(0, j).text = saved.columns[j]
# Add the body of the data frame
for i in range(saved.shape[0]):
    row = t.add_row()
    for j in range(saved.shape[1]):
        cell = saved.iat[i, j]
        row.cells[j].text = str(cell)
# Save the Word doc
doc.save('test_bal1.docx')


    


# %%            - IV - DOCUMENTATION & REFERENCES          –––––––

'''



    
    I.2.a - Retrieve data       
    
            How to remove first digit from a bigger number in python [closed]
            https://stackoverflow.com/questions/69858678/how-to-remove-first-digit-from-a-bigger-number-in-python
            
            How do I select rows from a DataFrame based on column values?
            https://stackoverflow.com/questions/17071871/how-do-i-select-rows-from-a-dataframe-based-on-column-values
            
            Pandas DataFrame loc[] Syntax and Examples
            https://sparkbyexamples.com/pandas/pandas-dataframe-loc/


    I.2.c - check parcellation process
    
            Extra ROI Variable in HO Atlas
            https://github.com/preprocessed-connectomes-project/abide/issues/9
            
            Is the number of ROIs for Harvard-Oxford atlas 110 or 111?
            https://www.nitrc.org/forum/message.php?msg_id=30606


––––––––––––––––––––––––––

I.3 - FUNCTION 1   
 
        Linear Regression in Python using Statsmodels
        https://datatofish.com/statsmodels-linear-regression/
        
        3 Ways to Convert String to Variable Name in Python
        2. Using locals() function to Convert a Python string to a Variable Name
        https://www.pythonpool.com/python-string-to-variable-name/      
        
        Delete Rows & Columns in DataFrames Quickly using Pandas Drop
        https://www.shanelynn.ie/pandas-drop-delete-dataframe-rows-columns/
        
        In sklearn regression, is there a command to return residuals for all records?
        https://stackoverflow.com/questions/55095437/in-sklearn-regression-is-there-a-command-to-return-residuals-for-all-records

        Unable to fit data due to "can't multiply sequence by non-int of type 'numpy.float64'" error
        https://stackoverflow.com/questions/55742096/unable-to-fit-data-due-to-cant-multiply-sequence-by-non-int-of-type-numpy-flo


––––––––––––––––––––––––––


I.4 - FUNCTION 2 

        Python create a dictionary in the loop
        https://tutorial.eyehunts.com/python/python-create-a-dictionary-in-the-loop-example-code/
                   
        
––––––––––––––––––––––––––

    II.4.a - raincloud plot function
    
            ................. code used to make the function .................
            Getting started with Raincloud plots in Python
            https://medium.com/mlearning-ai/getting-started-with-raincloud-plots-in-python-2ea5c2d01c11
            ..................................................................

––––––––––––––––––––––––––

    III.1.a - scanning site matrix
        
            How to sort dataframe based on a column in another dataframe in Pandas?
            https://stackoverflow.com/questions/45576800/how-to-sort-dataframe-based-on-a-column-in-another-dataframe-in-pandas

    
    III.1.b - plot scanning site mFC
    
            Merge, join, concatenate and compare
            https://pandas.pydata.org/docs/user_guide/merging.html
            
            seaborn.violinplot
            https://seaborn.pydata.org/generated/seaborn.violinplot.html
            
            Rotating axis labels in matplotlib and seaborn
            https://drawingfromdata.com/seaborn/matplotlib/visualization/rotate-axis-labels-matplotlib-seaborn.html
            

'''
